import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Length
from datetime import date
from num2words import num2words
from pathlib import Path

today = date.today().strftime('%d-%m-%Y')
dia = date.today().strftime('%d')
mes = date.today().strftime('%m')
ano = date.today().strftime('%Y')
if mes == '01':
    mesescrito = 'jAneiro'
elif mes == '02':
    mesescrito = 'fevereiro'
elif mes == '03':
    mesescrito = 'março'
elif mes == '04':
    mesescrito = 'abril'
elif mes == '05':
    mesescrito = 'maio'
elif mes == '06':
    mesescrito = 'junho'
elif mes == '07':
    mesescrito = 'julho'
elif mes == '08':
    mesescrito = 'agosto'
elif mes == '09':
    mesescrito = 'setembro'
elif mes == '10':
    mesescrito = 'outubro'
elif mes == '11':
    mesescrito = 'novembro'
elif mes == '12':
    mesescrito = 'dezembro'

root = tk.Tk()
root.title("ROCHA ENGENHARIA")
root.geometry("275x275")
root.iconbitmap('image/logo.ico')
root.focus()


# Fechar Pai
def on_closing():
    root.destroy()


# Janela da Obra
def Next():
    global janela2
    janela2 = tk.Toplevel()
    janela2.title('Cadastro Obra')
    janela2.iconbitmap('image/logo.ico')
    janela2.focus()
    root.withdraw()
    janela2.configure(bg='white')

    janela2.protocol("WM_DELETE_WINDOW", on_closing)

    label_nome = tk.Label(janela2, text="Dados da Obra",
                          font=("Arial", "11", "bold"))
    label_nome.grid(row=0, column=1)

    EndObraLabel = tk.Label(janela2, text="Endereço da obra")
    EndObraLabel.grid(row=1, column=1)

    global endobra
    endobra = tk.Entry(janela2)
    endobra.focus()
    endobra.bind("<Up>", lambda e: cb_cliente.focus())
    endobra.bind("<Down>", lambda e: NumObra.focus())
    endobra.bind("<Return>", lambda e: NumObra.focus())
    endobra["width"] = 30
    endobra.grid(row=1, column=2)

    NumObraLabel = tk.Label(janela2, text="Numero da Obra")
    NumObraLabel.grid(row=2, column=1)

    global NumObra
    NumObra = tk.Entry(janela2)
    NumObra.bind("<Up>", lambda e: endobra.focus())
    NumObra.bind("<Down>", lambda e: BairroObra.focus())
    NumObra.bind("<Return>", lambda e: BairroObra.focus())
    NumObra.bind("<KeyRelease>", format_numero1)
    NumObra["width"] = 30
    NumObra.grid(row=2, column=2)

    BairroObraLabel = tk.Label(janela2, text="Bairro da Obra")
    BairroObraLabel.grid(row=3, column=1)

    global BairroObra
    BairroObra = tk.Entry(janela2)
    BairroObra.bind("<Up>", lambda e: NumObra.focus())
    BairroObra.bind("<Down>", lambda e: LoteObra.focus())
    BairroObra.bind("<Return>", lambda e: LoteObra.focus())
    BairroObra["width"] = 30
    BairroObra.grid(row=3, column=2)

    LoteObraLabel = tk.Label(janela2, text="Lote da Obra")
    LoteObraLabel.grid(row=4, column=1)

    global LoteObra
    LoteObra = tk.Entry(janela2)
    LoteObra.bind("<Up>", lambda e: BairroObra.focus())
    LoteObra.bind("<Down>", lambda e: QuadraObra.focus())
    LoteObra.bind("<Return>", lambda e: QuadraObra.focus())
    LoteObra["width"] = 30
    LoteObra.grid(row=4, column=2)

    QuadraObraLabel = tk.Label(janela2, text="Quadra da obra")
    QuadraObraLabel.grid(row=5, column=1)

    global QuadraObra
    QuadraObra = tk.Entry(janela2)
    QuadraObra.bind("<Up>", lambda e: LoteObra.focus())
    QuadraObra.bind("<Down>", lambda e: CidadeObra.focus())
    QuadraObra.bind("<Return>", lambda e: CidadeObra.focus())
    QuadraObra["width"] = 30
    QuadraObra.grid(row=5, column=2)

    CidadeObraLabel = tk.Label(janela2, text="Cidade da obra")
    CidadeObraLabel.grid(row=6, column=1)

    listCidade = ["Americana", "Arthur Nogueira", "Bom Jesus dos Perdões", "Campinas", "Hortolândia", "Monte Mor",
                  "Nova Odessa", "Paulinia", "São Bernardo do Campo", "Sumaré"]

    global CidadeObra
    CidadeObra = ttk.Combobox(janela2, values=listCidade)
    CidadeObra.set("Hortolândia")
    CidadeObra.bind("<Up>", lambda e: QuadraObra.focus())
    CidadeObra.bind("<Down>", lambda e: QuarteiraoObra.focus())
    CidadeObra.bind("<Return>", lambda e: QuarteiraoObra.focus())
    CidadeObra["width"] = 27
    CidadeObra.grid(row=6, column=2)

    QuarteiraoObraLabel = tk.Label(janela2, text="Quarteirão da Obra")
    QuarteiraoObraLabel.grid(row=7, column=1)

    global QuarteiraoObra
    QuarteiraoObra = tk.Entry(janela2)
    QuarteiraoObra.bind("<Up>", lambda e: CidadeObra.focus())
    QuarteiraoObra.bind("<Down>", lambda e: TipoObra.focus())
    QuarteiraoObra.bind("<Return>", lambda e: TipoObra.focus())
    QuarteiraoObra["width"] = 30
    QuarteiraoObra.grid(row=7, column=2)

    TipoObraLabel = tk.Label(janela2, text="Tipo da Obra")
    TipoObraLabel.grid(row=8, column=1)

    listTipoObra = ["AVCB", "Comercial", "Habite-se", "Regularização", "Residencial", "Subdivisão"]

    global TipoObra
    TipoObra = ttk.Combobox(janela2, values=listTipoObra)
    TipoObra.set("Residencial")
    TipoObra.bind("<Up>", lambda e: QuarteiraoObra.focus())
    TipoObra.bind("<Down>", lambda e: AreaObra.focus())
    TipoObra.bind("<Return>", lambda e: AreaObra.focus())
    TipoObra["width"] = 27
    TipoObra.grid(row=8, column=2)

    AreaObraLabel = tk.Label(janela2, text="Area da Obra")
    AreaObraLabel.grid(row=9, column=1)

    global AreaObra
    AreaObra = tk.Entry(janela2)
    AreaObra.bind("<Up>", lambda e: TipoObra.focus())
    AreaObra.bind("<Down>", lambda e: ArtObra.focus())
    AreaObra.bind("<Return>", lambda e: ArtObra.focus())
    AreaObra["width"] = 30
    AreaObra.grid(row=9, column=2)

    ArtObraLabel = tk.Label(janela2, text="ART da Obra")
    ArtObraLabel.grid(row=10, column=1)

    global ArtObra
    ArtObra = tk.Entry(janela2)
    ArtObra.bind("<Up>", lambda e: AreaObra.focus())
    ArtObra.bind("<Down>", lambda e: ValorContrato.focus())
    ArtObra.bind("<Return>", lambda e: ValorContrato.focus())
    ArtObra["width"] = 30
    ArtObra.grid(row=10, column=2)

    ValorContratoLabel = tk.Label(janela2, text="Valor da Obra")
    ValorContratoLabel.grid(row=11, column=1)

    global ValorContrato
    ValorContrato = tk.Entry(janela2)
    ValorContrato.bind("<Up>", lambda e: ArtObra.focus())
    ValorContrato.bind("<Down>", lambda e: ParcelContrato.focus())
    ValorContrato.bind("<Return>", lambda e: ValorContratoextenso.focus())
    ValorContrato.bind("<KeyRelease>", lambda event: extenso1())
    ValorContrato["width"] = 30
    ValorContrato.grid(row=11, column=2)

    ValorContratoextensoLabel = tk.Label(janela2, text="Valor da Obra Extenso")
    ValorContratoextensoLabel.grid(row=12, column=1)

    global ValorContratoextenso
    ValorContratoextenso = tk.Entry(janela2)
    ValorContratoextenso.bind("<Up>", lambda e: ValorContrato.focus())
    ValorContratoextenso.bind("<Down>", lambda e: ParcelContrato.focus())
    ValorContratoextenso.bind("<Return>", lambda e: ParcelContrato.focus())
    ValorContratoextenso["width"] = 30
    ValorContratoextenso.grid(row=12, column=2)

    ParcelContratoLabel = tk.Label(janela2, text="Quantidade de parcela da obra")
    ParcelContratoLabel.grid(row=13, column=1)

    global ParcelContrato
    ParcelContrato = tk.Entry(janela2)
    ParcelContrato.bind("<Up>", lambda e: ValorContrato.focus())
    ParcelContrato.bind("<Down>", lambda e: ParcelContratoextenso.focus())
    ParcelContrato.bind("<Return>", lambda e: ParcelContratoextenso.focus())
    ParcelContrato.bind("<KeyRelease>", lambda event: extenso2())
    ParcelContrato["width"] = 30
    ParcelContrato.grid(row=13, column=2)

    ParcelContratoextensoLabel = tk.Label(janela2, text="Quantidade de parcela da obra Extenso")
    ParcelContratoextensoLabel.grid(row=14, column=1)

    global ParcelContratoextenso
    ParcelContratoextenso = tk.Entry(janela2)
    ParcelContratoextenso.bind("<Up>", lambda e: ParcelContrato.focus())
    ParcelContratoextenso.bind("<Down>", lambda e: ValorParcelaContrato.focus())
    ParcelContratoextenso.bind("<Return>", lambda e: ValorParcelaContrato.focus())
    ParcelContratoextenso["width"] = 30
    ParcelContratoextenso.grid(row=14, column=2)

    ValorParcelaContratoLabel = tk.Label(janela2, text="Valor da(s) parcela(s) da obra")
    ValorParcelaContratoLabel.grid(row=15, column=1)

    global ValorParcelaContrato
    ValorParcelaContrato = tk.Entry(janela2)
    ValorParcelaContrato.bind("<Up>", lambda e: ParcelContratoextenso.focus())
    ValorParcelaContrato.bind("<Down>", lambda e: ValorParcelaContratoextenso.focus())
    ValorParcelaContrato.bind("<Return>", lambda e: ValorParcelaContratoextenso.focus())
    ValorParcelaContrato.bind("<KeyRelease>", lambda event: extenso3())
    ValorParcelaContrato["width"] = 30
    ValorParcelaContrato.grid(row=15, column=2)

    ParcelContratoextensoLabel = tk.Label(janela2, text="Quantidade de parcela da obra Extenso")
    ParcelContratoextensoLabel.grid(row=16, column=1)

    global ValorParcelaContratoextenso
    ValorParcelaContratoextenso = tk.Entry(janela2)
    ValorParcelaContratoextenso.bind("<Up>", lambda e: ValorParcelaContrato.focus())
    ValorParcelaContratoextenso.bind("<Down>", lambda e: DiaVencimento.focus())
    ValorParcelaContratoextenso.bind("<Return>", lambda e: DiaVencimento.focus())
    ValorParcelaContratoextenso["width"] = 30
    ValorParcelaContratoextenso.grid(row=16, column=2)

    DiaVencimentoLabel = tk.Label(janela2, text="Dia de vencimento do Contrato")
    DiaVencimentoLabel.grid(row=17, column=1)

    listDiaVencimento = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12", "13", "14", "15", "16",
                         "17", "18", "19", "20", "21", "22", "23", "24", "25", "26", "27", "28", "29", "30", "31"]

    global DiaVencimento
    DiaVencimento = ttk.Combobox(janela2, values=listDiaVencimento)
    DiaVencimento.set("15")
    DiaVencimento.bind("<Up>", lambda e: ValorParcelaContratoextenso.focus())
    DiaVencimento.bind("<Down>", lambda e: InicioContrato.focus())
    DiaVencimento.bind("<Return>", lambda e: InicioContrato.focus())
    DiaVencimento["width"] = 27
    DiaVencimento.grid(row=17, column=2)

    InicioContratoLabel = tk.Label(janela2, text="Data do inicio do contrato")
    InicioContratoLabel.grid(row=18, column=1)

    global InicioContrato
    InicioContrato = tk.Entry(janela2)
    InicioContrato.bind("<Up>", lambda e: DiaVencimento.focus())
    InicioContrato.bind("<Down>", lambda e: Visita.focus())
    InicioContrato.bind("<Return>", lambda e: Visita.focus())
    InicioContrato.bind("<KeyRelease>", format_date1)
    InicioContrato["width"] = 30
    InicioContrato.grid(row=18, column=2)

    VisitaLabel = tk.Label(janela2, text="Valor Visita Técnica")
    VisitaLabel.grid(row=19, column=1)

    global Visita
    Visita = tk.Entry(janela2)
    Visita.bind("<Up>", lambda e: InicioContrato.focus())
    Visita.bind("<Down>", lambda e: cb_cliente.focus())
    Visita.bind("<Return>", lambda e: cb_cliente.focus())
    Visita.bind("<KeyRelease>", lambda event: extenso4())
    Visita["width"] = 30
    Visita.grid(row=19, column=2)

    VisitaextensoLabel = tk.Label(janela2, text="Valor Visita Técnica por Extenso")
    VisitaextensoLabel.grid(row=20, column=1)

    global Visitaextenso
    Visitaextenso = tk.Entry(janela2)
    Visitaextenso.bind("<Up>", lambda e: Visita.focus())
    Visitaextenso.bind("<Down>", lambda e: cb_cliente.focus())
    Visitaextenso.bind("<Return>", lambda e: cb_cliente.focus())
    Visitaextenso["width"] = 30
    Visitaextenso.grid(row=20, column=2)

    lb_Cliente = tk.Label(janela2, text="Quantidade de Clientes")
    lb_Cliente.grid(row=21, column=1)

    listCLiente = ["Um Cliente", "Dois Clientes", "Três Clientes", "Quatro Clientes"]

    global cb_cliente
    cb_cliente = ttk.Combobox(janela2, values=listCLiente)
    cb_cliente.set("Um Cliente")
    cb_cliente.bind("<KeyPress>", lambda e: on_button() if e.char == '\r' else None)
    cb_cliente.bind("<Up>", lambda e: Visita.focus())
    cb_cliente.bind("<Down>", lambda e: endobra.focus())
    cb_cliente.bind("<Return>", lambda e: on_button())
    cb_cliente["width"] = 27
    cb_cliente.grid(row=21, column=2)

    bntMenu = tk.Button(janela2, text="Voltar Menu")
    bntMenu["command"] = voltamenu
    bntMenu.grid(row=22, column=1)

    global bntCliente
    bntCliente = tk.Button(janela2, text="Dados CLiente")
    bntCliente["command"] = on_button
    bntCliente.grid(row=22, column=2)


# Comando para voltar para o main principal
def voltamenu():
    MsgBox = tk.messagebox.askquestion("ATENÇÃO", "Tem certeza que quer voltar? Se voltar perderá todos os dados",
                                       icon='warning')
    if MsgBox == 'yes':
        root.deiconify()
        janela2.destroy()
        JanelaCliente1.destroy()
        JanelaCliente2.destroy()
        JanelaCliente3.destroy()


# Comando para ir na janela de cliente
def on_button():
    if cb_cliente.get() == "Um Cliente":
        bntCliente["command"] = clienteone()



    elif cb_cliente.get() == "Dois Clientes":
        bntCliente["command"] = clientetwo()


    elif cb_cliente.get() == "Três Clientes":
        bntCliente["command"] = clientethree()

    elif cb_cliente.get() == "Quatro Clientes":
        bntCliente["command"] = clientefour()


# Janela para um cliente
def clienteone():
    global JanelaCliente1
    JanelaCliente1 = tk.Toplevel()
    JanelaCliente1.title('Cadastro Cliente')
    JanelaCliente1.iconbitmap('image/logo.ico')
    JanelaCliente1.focus()
    janela2.withdraw()
    JanelaCliente1.configure(bg='white')

    JanelaCliente1.protocol("WM_DELETE_WINDOW", on_closing)

    label_nome = tk.Label(JanelaCliente1, text="Dados do CLiente",
                          font=("Arial", "11", "bold"))
    label_nome.grid(row=0, column=1)

    NomeCLiente1Label = tk.Label(JanelaCliente1, text="Nome do CLiente")
    NomeCLiente1Label.grid(row=1, column=1)

    global NomeCliente1
    NomeCliente1 = tk.Entry(JanelaCliente1)
    NomeCliente1.focus()
    NomeCliente1.bind("<Up>", lambda e: CelCliente1.focus())
    NomeCliente1.bind("<Down>", lambda e: CPFCliente1.focus())
    NomeCliente1.bind("<Return>", lambda e: CPFCliente1.focus())
    NomeCliente1["width"] = 30
    NomeCliente1.grid(row=1, column=2)

    CPFCliente1Label = tk.Label(JanelaCliente1, text="CPF do Cliente")
    CPFCliente1Label.grid(row=2, column=1)

    global CPFCliente1
    CPFCliente1 = tk.Entry(JanelaCliente1)
    CPFCliente1.bind("<Up>", lambda e: NomeCliente1.focus())
    CPFCliente1.bind("<Down>", lambda e: RGCliente1.focus())
    CPFCliente1.bind("<Return>", lambda e: RGCliente1.focus())
    CPFCliente1.bind("<KeyRelease>", format_cpf1)
    CPFCliente1["width"] = 30
    CPFCliente1.grid(row=2, column=2)

    RGCliente1Label = tk.Label(JanelaCliente1, text="RG do Cliente")
    RGCliente1Label.grid(row=3, column=1)

    global RGCliente1
    RGCliente1 = tk.Entry(JanelaCliente1)
    RGCliente1.bind("<Up>", lambda e: CPFCliente1.focus())
    RGCliente1.bind("<Down>", lambda e: EndCliente1.focus())
    RGCliente1.bind("<Return>", lambda e: EndCliente1.focus())
    RGCliente1.bind("<KeyRelease>", format_rg1)
    RGCliente1["width"] = 30
    RGCliente1.grid(row=3, column=2)

    EndCliente1Label = tk.Label(JanelaCliente1, text="Endereço do Cliente")
    EndCliente1Label.grid(row=4, column=1)

    global EndCliente1
    EndCliente1 = tk.Entry(JanelaCliente1)
    EndCliente1.bind("<Up>", lambda e: RGCliente1.focus())
    EndCliente1.bind("<Down>", lambda e: NCliente1.focus())
    EndCliente1.bind("<Return>", lambda e: NCliente1.focus())
    EndCliente1["width"] = 30
    EndCliente1.grid(row=4, column=2)

    NCliente1Label = tk.Label(JanelaCliente1, text="Numero do Cliente")
    NCliente1Label.grid(row=5, column=1)

    global NCliente1
    NCliente1 = tk.Entry(JanelaCliente1)
    NCliente1.bind("<Up>", lambda e: EndCliente1.focus())
    NCliente1.bind("<Down>", lambda e: BairroCliente1.focus())
    NCliente1.bind("<Return>", lambda e: BairroCliente1.focus())
    NCliente1["width"] = 30
    NCliente1.grid(row=5, column=2)

    BairroCliente1Label = tk.Label(JanelaCliente1, text="Bairro do Cliente")
    BairroCliente1Label.grid(row=6, column=1)

    global BairroCliente1
    BairroCliente1 = tk.Entry(JanelaCliente1)
    BairroCliente1.bind("<Up>", lambda e: NCliente1.focus())
    BairroCliente1.bind("<Down>", lambda e: CidadeCliente1.focus())
    BairroCliente1.bind("<Return>", lambda e: CidadeCliente1.focus())
    BairroCliente1["width"] = 30
    BairroCliente1.grid(row=6, column=2)

    CidadeCliente1Label = tk.Label(JanelaCliente1, text="Cidade do Cliente")
    CidadeCliente1Label.grid(row=7, column=1)

    listCidade = ["Americana", "Arthur Nogueira", "Bom Jesus dos Perdões", "Campinas", "Hortolândia", "Monte Mor",
                  "Nova Odessa", "Paulinia", "São Bernardo do Campo", "Sumaré"]

    global CidadeCliente1
    CidadeCliente1 = ttk.Combobox(JanelaCliente1, values=listCidade)
    CidadeCliente1.set("Hortolândia")
    CidadeCliente1.bind("<Up>", lambda e: BairroCliente1.focus())
    CidadeCliente1.bind("<Down>", lambda e: EstadoCliente1.focus())
    CidadeCliente1.bind("<Return>", lambda e: EstadoCliente1.focus())
    CidadeCliente1["width"] = 27
    CidadeCliente1.grid(row=7, column=2)

    EstadoCliente1Label = tk.Label(JanelaCliente1, text="Estado do CLiente")
    EstadoCliente1Label.grid(row=8, column=1)

    listEstadoCliente1 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente1
    EstadoCliente1 = ttk.Combobox(JanelaCliente1, values=listEstadoCliente1)
    EstadoCliente1.bind("<Up>", lambda e: BairroCliente1.focus())
    EstadoCliente1.bind("<Down>", lambda e: CEPCliente1.focus())
    EstadoCliente1.bind("<Return>", lambda e: CEPCliente1.focus())
    EstadoCliente1.set("SP")
    EstadoCliente1["width"] = 27
    EstadoCliente1.grid(row=8, column=2)

    CEPCliente1Label = tk.Label(JanelaCliente1, text="CEP do Cliente")
    CEPCliente1Label.grid(row=9, column=1)

    global CEPCliente1
    CEPCliente1 = tk.Entry(JanelaCliente1)
    CEPCliente1.bind("<Up>", lambda e: EstadoCliente1.focus())
    CEPCliente1.bind("<Down>", lambda e: NacionalidadeCliente1.focus())
    CEPCliente1.bind("<Return>", lambda e: NacionalidadeCliente1.focus())
    CEPCliente1.bind("<KeyRelease>", format_CEP1)
    CEPCliente1["width"] = 30
    CEPCliente1.grid(row=9, column=2)

    NacionalidadeCliente1Label = tk.Label(JanelaCliente1, text="Nacionalidade do Cliente")
    NacionalidadeCliente1Label.grid(row=10, column=1)

    global NacionalidadeCliente1
    NacionalidadeCliente1 = tk.Entry(JanelaCliente1)
    NacionalidadeCliente1.bind("<Up>", lambda e: CEPCliente1.focus())
    NacionalidadeCliente1.bind("<Down>", lambda e: cb_EstCivilCliente1.focus())
    NacionalidadeCliente1.bind("<Return>", lambda e: cb_EstCivilCliente1.focus())
    NacionalidadeCliente1["width"] = 30
    NacionalidadeCliente1.grid(row=10, column=2)

    EstCivilCliente1Label = tk.Label(JanelaCliente1, text="Estado Civil Cliente")
    EstCivilCliente1Label.grid(row=11, column=1)

    listEstCivilCliente1 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    global cb_EstCivilCliente1
    cb_EstCivilCliente1 = ttk.Combobox(JanelaCliente1, values=listEstCivilCliente1)
    cb_EstCivilCliente1.bind("<Up>", lambda e: NacionalidadeCliente1.focus())
    cb_EstCivilCliente1.bind("<Down>", lambda e: ProfissaoCliente1.focus())
    cb_EstCivilCliente1.bind("<Return>", lambda e: ProfissaoCliente1.focus())
    cb_EstCivilCliente1.set("Solteiro(a)")
    cb_EstCivilCliente1["width"] = 27
    cb_EstCivilCliente1.grid(row=11, column=2)

    ProfissaoCliente1Label = tk.Label(JanelaCliente1, text="Profissão do Cliente")
    ProfissaoCliente1Label.grid(row=12, column=1)

    global ProfissaoCliente1
    ProfissaoCliente1 = tk.Entry(JanelaCliente1)
    ProfissaoCliente1.bind("<Up>", lambda e: cb_EstCivilCliente1.focus())
    ProfissaoCliente1.bind("<Down>", lambda e: EmailCliente1.focus())
    ProfissaoCliente1.bind("<Return>", lambda e: EmailCliente1.focus())
    ProfissaoCliente1["width"] = 30
    ProfissaoCliente1.grid(row=12, column=2)

    EmailCliente1Label = tk.Label(JanelaCliente1, text="Email do Cliente")
    EmailCliente1Label.grid(row=13, column=1)

    global EmailCliente1
    EmailCliente1 = tk.Entry(JanelaCliente1)
    EmailCliente1.bind("<Up>", lambda e: ProfissaoCliente1.focus())
    EmailCliente1.bind("<Down>", lambda e: CelCliente1.focus())
    EmailCliente1.bind("<Return>", lambda e: CelCliente1.focus())
    EmailCliente1["width"] = 30
    EmailCliente1.grid(row=13, column=2)

    CelCliente1Label = tk.Label(JanelaCliente1, text="Celular do Cliente")
    CelCliente1Label.grid(row=14, column=1)

    global CelCliente1
    CelCliente1 = tk.Entry(JanelaCliente1)
    CelCliente1.bind("<Up>", lambda e: EmailCliente1.focus())
    CelCliente1.bind("<Down>", lambda e: NomeCliente1.focus())
    CelCliente1.bind("<Return>", lambda e: NomeCliente1.focus())
    CelCliente1.bind("<KeyRelease>", format_celular1)
    CelCliente1["width"] = 30
    CelCliente1.grid(row=14, column=2)

    bntVoltar = tk.Button(JanelaCliente1, text="Voltar", command=volta)
    bntVoltar.grid(row=16, column=1)

    bntCliente = tk.Button(JanelaCliente1, text="Confirmar")
    bntCliente["command"] = confirmar1
    bntCliente.grid(row=16, column=2)


# Janela para dois clientes
def clientetwo():
    global JanelaCliente2
    JanelaCliente2 = tk.Toplevel()
    JanelaCliente2.title('Cadastro Cliente')
    JanelaCliente2.iconbitmap('image/logo.ico')
    JanelaCliente2.configure(bg='white')

    janela2.withdraw()

    JanelaCliente2.protocol("WM_DELETE_WINDOW", on_closing)

    label_nome = tk.Label(JanelaCliente2, text="Dados do CLiente 1",
                          font=("Arial", "11", "bold"))
    label_nome.grid(row=0, column=1)

    NomeCLiente1Label = tk.Label(JanelaCliente2, text="Nome do Cliente ")
    NomeCLiente1Label.grid(row=1, column=1)

    global NomeCliente1
    NomeCliente1 = tk.Entry(JanelaCliente2)
    NomeCliente1.focus()
    NomeCliente1.bind("<Up>", lambda e: CelCliente2.focus())
    NomeCliente1.bind("<Down>", lambda e: CPFCliente1.focus())
    NomeCliente1.bind("<Return>", lambda e: CPFCliente1.focus())
    NomeCliente1["width"] = 30
    NomeCliente1.grid(row=1, column=2)

    CPFCliente1Label = tk.Label(JanelaCliente2, text="CPF do Cliente")
    CPFCliente1Label.grid(row=2, column=1)

    global CPFCliente1
    CPFCliente1 = tk.Entry(JanelaCliente2)
    CPFCliente1.bind("<Up>", lambda e: NomeCliente1.focus())
    CPFCliente1.bind("<Down>", lambda e: RGCliente1.focus())
    CPFCliente1.bind("<Return>", lambda e: RGCliente1.focus())
    CPFCliente1.bind("<KeyRelease>", format_cpf1)
    CPFCliente1["width"] = 30
    CPFCliente1.grid(row=2, column=2)

    RGCliente1Label = tk.Label(JanelaCliente2, text="RG do Cliente")
    RGCliente1Label.grid(row=3, column=1)

    global RGCliente1
    RGCliente1 = tk.Entry(JanelaCliente2)
    RGCliente1.bind("<Up>", lambda e: CPFCliente1.focus())
    RGCliente1.bind("<Down>", lambda e: EndCliente1.focus())
    RGCliente1.bind("<Return>", lambda e: EndCliente1.focus())
    RGCliente1.bind("<KeyRelease>", format_rg1)
    RGCliente1["width"] = 30
    RGCliente1.grid(row=3, column=2)

    EndCliente1Label = tk.Label(JanelaCliente2, text="Endereço do Cliente")
    EndCliente1Label.grid(row=4, column=1)

    global EndCliente1
    EndCliente1 = tk.Entry(JanelaCliente2)
    EndCliente1.bind("<Up>", lambda e: RGCliente1.focus())
    EndCliente1.bind("<Down>", lambda e: NCliente1.focus())
    EndCliente1.bind("<Return>", lambda e: NCliente1.focus())
    EndCliente1["width"] = 30
    EndCliente1.grid(row=4, column=2)

    NCliente1Label = tk.Label(JanelaCliente2, text="Numero do Cliente")
    NCliente1Label.grid(row=5, column=1)

    global NCliente1
    NCliente1 = tk.Entry(JanelaCliente2)
    NCliente1.bind("<Up>", lambda e: EndCliente1.focus())
    NCliente1.bind("<Down>", lambda e: BairroCliente1.focus())
    NCliente1.bind("<Return>", lambda e: BairroCliente1.focus())
    NCliente1["width"] = 30
    NCliente1.grid(row=5, column=2)

    BairroCliente1Label = tk.Label(JanelaCliente2, text="Bairro do Cliente")
    BairroCliente1Label.grid(row=6, column=1)

    global BairroCliente1
    BairroCliente1 = tk.Entry(JanelaCliente2)
    BairroCliente1.bind("<Up>", lambda e: NCliente1.focus())
    BairroCliente1.bind("<Down>", lambda e: CidadeCliente1.focus())
    BairroCliente1.bind("<Return>", lambda e: CidadeCliente1.focus())
    BairroCliente1["width"] = 30
    BairroCliente1.grid(row=6, column=2)

    CidadeCliente1Label = tk.Label(JanelaCliente2, text="Cidade do Cliente")
    CidadeCliente1Label.grid(row=7, column=1)

    listCidade = ["Americana", "Arthur Nogueira", "Bom Jesus dos Perdões", "Campinas", "Hortolândia", "Monte Mor",
                  "Nova Odessa", "Paulinia", "São Bernardo do Campo", "Sumaré"]

    global CidadeCliente1
    CidadeCliente1 = ttk.Combobox(JanelaCliente2, values=listCidade)
    CidadeCliente1.set("Hortolândia")
    CidadeCliente1.bind("<Up>", lambda e: BairroCliente1.focus())
    CidadeCliente1.bind("<Down>", lambda e: EstadoCliente1.focus())
    CidadeCliente1.bind("<Return>", lambda e: EstadoCliente1.focus())
    CidadeCliente1["width"] = 27
    CidadeCliente1.grid(row=7, column=2)

    EstadoCliente1Label = tk.Label(JanelaCliente2, text="Estado do CLiente")
    EstadoCliente1Label.grid(row=8, column=1)

    listEstadoCliente1 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente1
    EstadoCliente1 = ttk.Combobox(JanelaCliente2, values=listEstadoCliente1)
    EstadoCliente1.bind("<Up>", lambda e: CidadeCliente1.focus())
    EstadoCliente1.bind("<Down>", lambda e: CEPCliente1.focus())
    EstadoCliente1.bind("<Return>", lambda e: CEPCliente1.focus())
    EstadoCliente1.set("SP")
    EstadoCliente1["width"] = 27
    EstadoCliente1.grid(row=8, column=2)

    CEPCliente1Label = tk.Label(JanelaCliente2, text="CEP do Cliente")
    CEPCliente1Label.grid(row=9, column=1)

    global CEPCliente1
    CEPCliente1 = tk.Entry(JanelaCliente2)
    CEPCliente1.bind("<Up>", lambda e: EstadoCliente1.focus())
    CEPCliente1.bind("<Down>", lambda e: NacionalidadeCliente1.focus())
    CEPCliente1.bind("<Return>", lambda e: NacionalidadeCliente1.focus())
    CEPCliente1.bind("<KeyRelease>", format_CEP1)
    CEPCliente1["width"] = 30
    CEPCliente1.grid(row=9, column=2)

    NacionalidadeCliente1Label = tk.Label(JanelaCliente2, text="Nacionalidade do Cliente")
    NacionalidadeCliente1Label.grid(row=10, column=1)

    global NacionalidadeCliente1
    NacionalidadeCliente1 = tk.Entry(JanelaCliente2)
    NacionalidadeCliente1.bind("<Up>", lambda e: CEPCliente1.focus())
    NacionalidadeCliente1.bind("<Down>", lambda e: cb_EstCivilCliente1.focus())
    NacionalidadeCliente1.bind("<Return>", lambda e: cb_EstCivilCliente1.focus())
    NacionalidadeCliente1["width"] = 30
    NacionalidadeCliente1.grid(row=10, column=2)

    listEstCivilCliente1 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente1Label = tk.Label(JanelaCliente2, text="Estado Civil Cliente")
    EstCivilCliente1Label.grid(row=11, column=1)

    global cb_EstCivilCliente1
    cb_EstCivilCliente1 = ttk.Combobox(JanelaCliente2, values=listEstCivilCliente1)
    cb_EstCivilCliente1.bind("<Up>", lambda e: NacionalidadeCliente1.focus())
    cb_EstCivilCliente1.bind("<Down>", lambda e: ProfissaoCliente1.focus())
    cb_EstCivilCliente1.bind("<Return>", lambda e: ProfissaoCliente1.focus())
    cb_EstCivilCliente1.set("Solteiro(a)")
    cb_EstCivilCliente1["width"] = 27
    cb_EstCivilCliente1.grid(row=11, column=2)

    ProfissaoCliente1Label = tk.Label(JanelaCliente2, text="Profissão do Cliente")
    ProfissaoCliente1Label.grid(row=12, column=1)

    global ProfissaoCliente1
    ProfissaoCliente1 = tk.Entry(JanelaCliente2)
    ProfissaoCliente1.bind("<Up>", lambda e: cb_EstCivilCliente1.focus())
    ProfissaoCliente1.bind("<Down>", lambda e: EmailCliente1.focus())
    ProfissaoCliente1.bind("<Return>", lambda e: EmailCliente1.focus())
    ProfissaoCliente1["width"] = 30
    ProfissaoCliente1.grid(row=12, column=2)

    EmailCliente1Label = tk.Label(JanelaCliente2, text="Email do Cliente")
    EmailCliente1Label.grid(row=13, column=1)

    global EmailCliente1
    EmailCliente1 = tk.Entry(JanelaCliente2)
    EmailCliente1.bind("<Up>", lambda e: ProfissaoCliente1.focus())
    EmailCliente1.bind("<Down>", lambda e: CelCliente1.focus())
    EmailCliente1.bind("<Return>", lambda e: CelCliente1.focus())
    EmailCliente1["width"] = 30
    EmailCliente1.grid(row=13, column=2)

    CelCliente1Label = tk.Label(JanelaCliente2, text="Celular do Cliente")
    CelCliente1Label.grid(row=14, column=1)

    global CelCliente1
    CelCliente1 = tk.Entry(JanelaCliente2)
    CelCliente1.bind("<Up>", lambda e: EmailCliente1.focus())
    CelCliente1.bind("<Down>", lambda e: NomeCliente2.focus())
    CelCliente1.bind("<Return>", lambda e: NomeCliente2.focus())
    CelCliente1.bind("<KeyRelease>", format_celular1)
    CelCliente1["width"] = 30
    CelCliente1.grid(row=14, column=2)

    # --------------------------------------------------CLIENTE 2 ----------------------------------------------------------------------------------

    label_nome2 = tk.Label(JanelaCliente2, text="Dados do Cliente 2",
                           font=("Arial", "11", "bold"))
    label_nome2.grid(row=0, column=4)

    NomeCLiente2Label = tk.Label(JanelaCliente2, text="Nome do Cliente 2")
    NomeCLiente2Label.grid(row=1, column=4)

    global NomeCliente2
    NomeCliente2 = tk.Entry(JanelaCliente2)
    NomeCliente2.bind("<Up>", lambda e: CelCliente1.focus())
    NomeCliente2.bind("<Down>", lambda e: CPFCliente2.focus())
    NomeCliente2.bind("<Return>", lambda e: CPFCliente2.focus())
    NomeCliente2["width"] = 30
    NomeCliente2.grid(row=1, column=5)

    CPFCliente2Label = tk.Label(JanelaCliente2, text="CPF do Cliente 2")
    CPFCliente2Label.grid(row=2, column=4)

    global CPFCliente2
    CPFCliente2 = tk.Entry(JanelaCliente2)
    CPFCliente2.bind("<Up>", lambda e: NomeCliente2.focus())
    CPFCliente2.bind("<Down>", lambda e: RGCliente2.focus())
    CPFCliente2.bind("<Return>", lambda e: RGCliente2.focus())
    CPFCliente2.bind("<KeyRelease>", format_cpf2)
    CPFCliente2["width"] = 30
    CPFCliente2.grid(row=2, column=5)

    RGCliente2Label = tk.Label(JanelaCliente2, text="RG do Cliente")
    RGCliente2Label.grid(row=3, column=4)

    global RGCliente2
    RGCliente2 = tk.Entry(JanelaCliente2)
    RGCliente2.bind("<Up>", lambda e: CPFCliente2.focus())
    RGCliente2.bind("<Down>", lambda e: EndCliente2.focus())
    RGCliente2.bind("<Return>", lambda e: EndCliente2.focus())
    RGCliente2.bind("<KeyRelease>", format_rg2)
    RGCliente2["width"] = 30
    RGCliente2.grid(row=3, column=5)

    EndCliente2Label = tk.Label(JanelaCliente2, text="Endereço do Cliente")
    EndCliente2Label.grid(row=4, column=4)

    global EndCliente2
    EndCliente2 = tk.Entry(JanelaCliente2)
    EndCliente2.bind("<Up>", lambda e: RGCliente2.focus())
    EndCliente2.bind("<Down>", lambda e: NCliente2.focus())
    EndCliente2.bind("<Return>", lambda e: NCliente2.focus())
    EndCliente2["width"] = 30
    EndCliente2.grid(row=4, column=5)

    NCliente2Label = tk.Label(JanelaCliente2, text="Numero do Cliente")
    NCliente2Label.grid(row=5, column=4)

    global NCliente2
    NCliente2 = tk.Entry(JanelaCliente2)
    NCliente2.bind("<Up>", lambda e: EndCliente2.focus())
    NCliente2.bind("<Down>", lambda e: BairroCliente2.focus())
    NCliente2.bind("<Return>", lambda e: BairroCliente2.focus())
    NCliente2["width"] = 30
    NCliente2.grid(row=5, column=5)

    BairroCliente2Label = tk.Label(JanelaCliente2, text="Bairro do Cliente")
    BairroCliente2Label.grid(row=6, column=4)

    global BairroCliente2
    BairroCliente2 = tk.Entry(JanelaCliente2)
    BairroCliente2.bind("<Up>", lambda e: NCliente2.focus())
    BairroCliente2.bind("<Down>", lambda e: CidadeCliente2.focus())
    BairroCliente2.bind("<Return>", lambda e: CidadeCliente2.focus())
    BairroCliente2["width"] = 30
    BairroCliente2.grid(row=6, column=5)

    CidadeCliente2Label = tk.Label(JanelaCliente2, text="Cidade do Cliente")
    CidadeCliente2Label.grid(row=7, column=4)

    listCidade = ["Americana", "Arthur Nogueira", "Bom Jesus dos Perdões", "Campinas", "Hortolândia", "Monte Mor",
                  "Nova Odessa", "Paulinia", "São Bernardo do Campo", "Sumaré"]

    global CidadeCliente2
    CidadeCliente2 = ttk.Combobox(JanelaCliente2, values=listCidade)
    CidadeCliente2.set("Hortolândia")
    CidadeCliente2.bind("<Up>", lambda e: BairroCliente2.focus())
    CidadeCliente2.bind("<Down>", lambda e: EstadoCliente2.focus())
    CidadeCliente2.bind("<Return>", lambda e: EstadoCliente2.focus())
    CidadeCliente2["width"] = 27
    CidadeCliente2.grid(row=7, column=5)

    EstadoCliente2Label = tk.Label(JanelaCliente2, text="Estado do CLiente")
    EstadoCliente2Label.grid(row=8, column=4)

    listEstadoCliente2 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente2
    EstadoCliente2 = ttk.Combobox(JanelaCliente2, values=listEstadoCliente2)
    EstadoCliente2.bind("<Up>", lambda e: CidadeCliente2.focus())
    EstadoCliente2.bind("<Down>", lambda e: CEPCliente2.focus())
    EstadoCliente2.bind("<Return>", lambda e: CEPCliente2.focus())
    EstadoCliente2.set("SP")
    EstadoCliente2["width"] = 27
    EstadoCliente2.grid(row=8, column=5)

    CEPCliente2Label = tk.Label(JanelaCliente2, text="CEP do Cliente")
    CEPCliente2Label.grid(row=9, column=4)

    global CEPCliente2
    CEPCliente2 = tk.Entry(JanelaCliente2)
    CEPCliente2.bind("<Up>", lambda e: EstadoCliente2.focus())
    CEPCliente2.bind("<Down>", lambda e: NacionalidadeCliente2.focus())
    CEPCliente2.bind("<Return>", lambda e: NacionalidadeCliente2.focus())
    CEPCliente2.bind("<KeyRelease>", format_CEP2)
    CEPCliente2["width"] = 30
    CEPCliente2.grid(row=9, column=5)

    NacionalidadeCliente2Label = tk.Label(JanelaCliente2, text="Nacionalidade do Cliente")
    NacionalidadeCliente2Label.grid(row=10, column=4)

    global NacionalidadeCliente2
    NacionalidadeCliente2 = tk.Entry(JanelaCliente2)
    NacionalidadeCliente2.bind("<Up>", lambda e: CEPCliente2.focus())
    NacionalidadeCliente2.bind("<Down>", lambda e: cb_EstCivilCliente2.focus())
    NacionalidadeCliente2.bind("<Return>", lambda e: cb_EstCivilCliente2.focus())
    NacionalidadeCliente2["width"] = 30
    NacionalidadeCliente2.grid(row=10, column=5)

    listEstCivilCliente2 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente2Label = tk.Label(JanelaCliente2, text="Estado Civil Cliente")
    EstCivilCliente2Label.grid(row=11, column=4)

    global cb_EstCivilCliente2
    cb_EstCivilCliente2 = ttk.Combobox(JanelaCliente2, values=listEstCivilCliente2)
    cb_EstCivilCliente2.set("Solteiro(a)")
    cb_EstCivilCliente2.bind("<Up>", lambda e: NacionalidadeCliente2.focus())
    cb_EstCivilCliente2.bind("<Down>", lambda e: ProfissaoCliente2.focus())
    cb_EstCivilCliente2.bind("<Return>", lambda e: ProfissaoCliente2.focus())
    cb_EstCivilCliente2["width"] = 27
    cb_EstCivilCliente2.grid(row=11, column=5)

    ProfissaoCliente2Label = tk.Label(JanelaCliente2, text="Profissão do Cliente")
    ProfissaoCliente2Label.grid(row=12, column=4)

    global ProfissaoCliente2
    ProfissaoCliente2 = tk.Entry(JanelaCliente2)
    ProfissaoCliente2.bind("<Up>", lambda e: cb_EstCivilCliente2.focus())
    ProfissaoCliente2.bind("<Down>", lambda e: EmailCliente2.focus())
    ProfissaoCliente2.bind("<Return>", lambda e: EmailCliente2.focus())
    ProfissaoCliente2["width"] = 30
    ProfissaoCliente2.grid(row=12, column=5)

    EmailCliente2Label = tk.Label(JanelaCliente2, text="Email do Cliente")
    EmailCliente2Label.grid(row=13, column=4)

    global EmailCliente2
    EmailCliente2 = tk.Entry(JanelaCliente2)
    EmailCliente2.bind("<Up>", lambda e: ProfissaoCliente2.focus())
    EmailCliente2.bind("<Down>", lambda e: CelCliente2.focus())
    EmailCliente2.bind("<Return>", lambda e: CelCliente2.focus())
    EmailCliente2["width"] = 30
    EmailCliente2.grid(row=13, column=5)

    CelCliente2Label = tk.Label(JanelaCliente2, text="Celular do Cliente")
    CelCliente2Label.grid(row=14, column=4)

    global CelCliente2
    CelCliente2 = tk.Entry(JanelaCliente2)
    CelCliente2.bind("<Up>", lambda e: EmailCliente2.focus())
    CelCliente2.bind("<Down>", lambda e: NomeCliente1.focus())
    CelCliente2.bind("<Return>", lambda e: confirmar2())
    CelCliente2.bind("<KeyRelease>", format_celular2)
    CelCliente2["width"] = 30
    CelCliente2.grid(row=14, column=5)

    bntVoltar2 = tk.Button(JanelaCliente2, text="Voltar", command=volta2)
    bntVoltar2.grid(row=16, column=2)

    bntCliente2 = tk.Button(JanelaCliente2, text="Confirmar")
    bntCliente2["command"] = confirmar2
    bntCliente2.grid(row=16, column=4)

    bntCopiar = tk.Button(JanelaCliente2, text="Copiar")
    bntCopiar["command"] = copiar21
    bntCopiar.grid(row=16, column=5)


# Janela para três clientes
def clientethree():
    global JanelaCliente3
    JanelaCliente3 = tk.Toplevel()
    JanelaCliente3.title('Cadastro Cliente')
    JanelaCliente3.iconbitmap('image/logo.ico')
    JanelaCliente3.configure(bg='white')

    janela2.withdraw()

    JanelaCliente3.protocol("WM_DELETE_WINDOW", on_closing)

    label_nome = tk.Label(JanelaCliente3, text="Dados do CLiente 1",
                          font=("Arial", "11", "bold"))
    label_nome.grid(row=0, column=1)

    NomeCLiente1Label = tk.Label(JanelaCliente3, text="Nome do Cliente ")
    NomeCLiente1Label.grid(row=1, column=1)

    global NomeCliente1
    NomeCliente1 = tk.Entry(JanelaCliente3)
    NomeCliente1.focus()
    NomeCliente1.bind("<Up>", lambda e: CelCliente3.focus())
    NomeCliente1.bind("<Down>", lambda e: CPFCliente1.focus())
    NomeCliente1.bind("<Return>", lambda e: CPFCliente1.focus())
    NomeCliente1["width"] = 30
    NomeCliente1.grid(row=1, column=2)

    CPFCliente1Label = tk.Label(JanelaCliente3, text="CPF do Cliente")
    CPFCliente1Label.grid(row=2, column=1)

    global CPFCliente1
    CPFCliente1 = tk.Entry(JanelaCliente3)
    CPFCliente1.bind("<Up>", lambda e: NomeCliente1.focus())
    CPFCliente1.bind("<Down>", lambda e: RGCliente1.focus())
    CPFCliente1.bind("<Return>", lambda e: RGCliente1.focus())
    CPFCliente1.bind("<KeyRelease>", format_cpf1)
    CPFCliente1["width"] = 30
    CPFCliente1.grid(row=2, column=2)

    RGCliente1Label = tk.Label(JanelaCliente3, text="RG do Cliente")
    RGCliente1Label.grid(row=3, column=1)

    global RGCliente1
    RGCliente1 = tk.Entry(JanelaCliente3)
    RGCliente1.bind("<Up>", lambda e: CPFCliente1.focus())
    RGCliente1.bind("<Down>", lambda e: EndCliente1.focus())
    RGCliente1.bind("<Return>", lambda e: EndCliente1.focus())
    RGCliente1.bind("<KeyRelease>", format_rg1)
    RGCliente1["width"] = 30
    RGCliente1.grid(row=3, column=2)

    EndCliente1Label = tk.Label(JanelaCliente3, text="Endereço do Cliente")
    EndCliente1Label.grid(row=4, column=1)

    global EndCliente1
    EndCliente1 = tk.Entry(JanelaCliente3)
    EndCliente1.bind("<Up>", lambda e: RGCliente1.focus())
    EndCliente1.bind("<Down>", lambda e: NCliente1.focus())
    EndCliente1.bind("<Return>", lambda e: NCliente1.focus())
    EndCliente1["width"] = 30
    EndCliente1.grid(row=4, column=2)

    NCliente1Label = tk.Label(JanelaCliente3, text="Numero do Cliente")
    NCliente1Label.grid(row=5, column=1)

    global NCliente1
    NCliente1 = tk.Entry(JanelaCliente3)
    NCliente1.bind("<Up>", lambda e: EndCliente1.focus())
    NCliente1.bind("<Down>", lambda e: BairroCliente1.focus())
    NCliente1.bind("<Return>", lambda e: BairroCliente1.focus())
    NCliente1["width"] = 30
    NCliente1.grid(row=5, column=2)

    BairroCliente1Label = tk.Label(JanelaCliente3, text="Bairro do Cliente")
    BairroCliente1Label.grid(row=6, column=1)

    global BairroCliente1
    BairroCliente1 = tk.Entry(JanelaCliente3)
    BairroCliente1.bind("<Up>", lambda e: NCliente1.focus())
    BairroCliente1.bind("<Down>", lambda e: CidadeCliente1.focus())
    BairroCliente1.bind("<Return>", lambda e: CidadeCliente1.focus())
    BairroCliente1["width"] = 30
    BairroCliente1.grid(row=6, column=2)

    CidadeCliente1Label = tk.Label(JanelaCliente3, text="Cidade do Cliente")
    CidadeCliente1Label.grid(row=7, column=1)

    global CidadeCliente1
    CidadeCliente1 = tk.Entry(JanelaCliente3)
    CidadeCliente1.bind("<Up>", lambda e: BairroCliente1.focus())
    CidadeCliente1.bind("<Down>", lambda e: EstadoCliente1.focus())
    CidadeCliente1.bind("<Return>", lambda e: EstadoCliente1.focus())
    CidadeCliente1["width"] = 30
    CidadeCliente1.grid(row=7, column=2)

    EstadoCliente1Label = tk.Label(JanelaCliente3, text="Estado do CLiente")
    EstadoCliente1Label.grid(row=8, column=1)

    listEstadoCliente1 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente1
    EstadoCliente1 = ttk.Combobox(JanelaCliente3, values=listEstadoCliente1)
    EstadoCliente1.bind("<Up>", lambda e: CidadeCliente1.focus())
    EstadoCliente1.bind("<Down>", lambda e: CEPCliente1.focus())
    EstadoCliente1.bind("<Return>", lambda e: CEPCliente1.focus())
    EstadoCliente1.set("SP")
    EstadoCliente1["width"] = 27
    EstadoCliente1.grid(row=8, column=2)

    CEPCliente1Label = tk.Label(JanelaCliente3, text="CEP do Cliente")
    CEPCliente1Label.grid(row=9, column=1)

    global CEPCliente1
    CEPCliente1 = tk.Entry(JanelaCliente3)
    CEPCliente1.bind("<Up>", lambda e: EstadoCliente1.focus())
    CEPCliente1.bind("<Down>", lambda e: NacionalidadeCliente1.focus())
    CEPCliente1.bind("<Return>", lambda e: NacionalidadeCliente1.focus())
    CEPCliente1.bind("<KeyRelease>", format_CEP1)
    CEPCliente1["width"] = 30
    CEPCliente1.grid(row=9, column=2)

    NacionalidadeCliente1Label = tk.Label(JanelaCliente3, text="Nacionalidade do Cliente")
    NacionalidadeCliente1Label.grid(row=10, column=1)

    global NacionalidadeCliente1
    NacionalidadeCliente1 = tk.Entry(JanelaCliente3)
    NacionalidadeCliente1.bind("<Up>", lambda e: CEPCliente1.focus())
    NacionalidadeCliente1.bind("<Down>", lambda e: cb_EstCivilCliente1.focus())
    NacionalidadeCliente1.bind("<Return>", lambda e: cb_EstCivilCliente1.focus())
    NacionalidadeCliente1["width"] = 30
    NacionalidadeCliente1.grid(row=10, column=2)

    listEstCivilCliente1 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente1Label = tk.Label(JanelaCliente3, text="Estado Civil Cliente")
    EstCivilCliente1Label.grid(row=11, column=1)

    global cb_EstCivilCliente1
    cb_EstCivilCliente1 = ttk.Combobox(JanelaCliente3, values=listEstCivilCliente1)
    cb_EstCivilCliente1.bind("<Up>", lambda e: NacionalidadeCliente1.focus())
    cb_EstCivilCliente1.bind("<Down>", lambda e: ProfissaoCliente1.focus())
    cb_EstCivilCliente1.bind("<Return>", lambda e: ProfissaoCliente1.focus())
    cb_EstCivilCliente1.set("Solteiro(a)")
    cb_EstCivilCliente1["width"] = 27
    cb_EstCivilCliente1.grid(row=11, column=2)

    ProfissaoCliente1Label = tk.Label(JanelaCliente3, text="Profissão do Cliente")
    ProfissaoCliente1Label.grid(row=12, column=1)

    global ProfissaoCliente1
    ProfissaoCliente1 = tk.Entry(JanelaCliente3)
    ProfissaoCliente1.bind("<Up>", lambda e: cb_EstCivilCliente1.focus())
    ProfissaoCliente1.bind("<Down>", lambda e: EmailCliente1.focus())
    ProfissaoCliente1.bind("<Return>", lambda e: EmailCliente1.focus())
    ProfissaoCliente1["width"] = 30
    ProfissaoCliente1.grid(row=12, column=2)

    EmailCliente1Label = tk.Label(JanelaCliente3, text="Email do Cliente")
    EmailCliente1Label.grid(row=13, column=1)

    global EmailCliente1
    EmailCliente1 = tk.Entry(JanelaCliente3)
    EmailCliente1.bind("<Up>", lambda e: ProfissaoCliente1.focus())
    EmailCliente1.bind("<Down>", lambda e: CelCliente1.focus())
    EmailCliente1.bind("<Return>", lambda e: CelCliente1.focus())
    EmailCliente1["width"] = 30
    EmailCliente1.grid(row=13, column=2)

    CelCliente1Label = tk.Label(JanelaCliente3, text="Celular do Cliente")
    CelCliente1Label.grid(row=14, column=1)

    global CelCliente1
    CelCliente1 = tk.Entry(JanelaCliente3)
    CelCliente1.bind("<Up>", lambda e: EmailCliente1.focus())
    CelCliente1.bind("<Down>", lambda e: NomeCliente2.focus())
    CelCliente1.bind("<Return>", lambda e: NomeCliente2.focus())
    CelCliente1.bind("<KeyRelease>", format_celular1)
    CelCliente1["width"] = 30
    CelCliente1.grid(row=14, column=2)

    # --------------------------------------------------CLIENTE 2 ----------------------------------------------------------------------------------

    label_nome2 = tk.Label(JanelaCliente3, text="Dados do Cliente 2",
                           font=("Arial", "11", "bold"))
    label_nome2.grid(row=0, column=4)

    NomeCLiente2Label = tk.Label(JanelaCliente3, text="Nome do Cliente 2")
    NomeCLiente2Label.grid(row=1, column=4)

    global NomeCliente2
    NomeCliente2 = tk.Entry(JanelaCliente3)
    NomeCliente2.bind("<Up>", lambda e: CelCliente1.focus())
    NomeCliente2.bind("<Down>", lambda e: CPFCliente2.focus())
    NomeCliente2.bind("<Return>", lambda e: CPFCliente2.focus())
    NomeCliente2["width"] = 30
    NomeCliente2.grid(row=1, column=5)

    CPFCliente2Label = tk.Label(JanelaCliente3, text="CPF do Cliente 2")
    CPFCliente2Label.grid(row=2, column=4)

    global CPFCliente2
    CPFCliente2 = tk.Entry(JanelaCliente3)
    CPFCliente2.bind("<Up>", lambda e: NomeCliente2.focus())
    CPFCliente2.bind("<Down>", lambda e: RGCliente2.focus())
    CPFCliente2.bind("<Return>", lambda e: RGCliente2.focus())
    CPFCliente2.bind("<KeyRelease>", format_cpf2)
    CPFCliente2["width"] = 30
    CPFCliente2.grid(row=2, column=5)

    RGCliente2Label = tk.Label(JanelaCliente3, text="RG do Cliente")
    RGCliente2Label.grid(row=3, column=4)

    global RGCliente2
    RGCliente2 = tk.Entry(JanelaCliente3)
    RGCliente2.bind("<Up>", lambda e: CPFCliente2.focus())
    RGCliente2.bind("<Down>", lambda e: EndCliente2.focus())
    RGCliente2.bind("<Return>", lambda e: EndCliente2.focus())
    RGCliente2.bind("<KeyRelease>", format_rg2)
    RGCliente2["width"] = 30
    RGCliente2.grid(row=3, column=5)

    EndCliente2Label = tk.Label(JanelaCliente3, text="Endereço do Cliente")
    EndCliente2Label.grid(row=4, column=4)

    global EndCliente2
    EndCliente2 = tk.Entry(JanelaCliente3)
    EndCliente2.bind("<Up>", lambda e: RGCliente2.focus())
    EndCliente2.bind("<Down>", lambda e: NCliente2.focus())
    EndCliente2.bind("<Return>", lambda e: NCliente2.focus())
    EndCliente2["width"] = 30
    EndCliente2.grid(row=4, column=5)

    NCliente2Label = tk.Label(JanelaCliente3, text="Numero do Cliente")
    NCliente2Label.grid(row=5, column=4)

    global NCliente2
    NCliente2 = tk.Entry(JanelaCliente3)
    NCliente2.bind("<Up>", lambda e: EndCliente2.focus())
    NCliente2.bind("<Down>", lambda e: BairroCliente2.focus())
    NCliente2.bind("<Return>", lambda e: BairroCliente2.focus())
    NCliente2["width"] = 30
    NCliente2.grid(row=5, column=5)

    BairroCliente2Label = tk.Label(JanelaCliente3, text="Bairro do Cliente")
    BairroCliente2Label.grid(row=6, column=4)

    global BairroCliente2
    BairroCliente2 = tk.Entry(JanelaCliente3)
    BairroCliente2.bind("<Up>", lambda e: NCliente2.focus())
    BairroCliente2.bind("<Down>", lambda e: CidadeCliente2.focus())
    BairroCliente2.bind("<Return>", lambda e: CidadeCliente2.focus())
    BairroCliente2["width"] = 30
    BairroCliente2.grid(row=6, column=5)

    CidadeCliente2Label = tk.Label(JanelaCliente3, text="Cidade do Cliente")
    CidadeCliente2Label.grid(row=7, column=4)

    global CidadeCliente2
    CidadeCliente2 = tk.Entry(JanelaCliente3)
    CidadeCliente2.bind("<Up>", lambda e: BairroCliente2.focus())
    CidadeCliente2.bind("<Down>", lambda e: EstadoCliente2.focus())
    CidadeCliente2.bind("<Return>", lambda e: EstadoCliente2.focus())
    CidadeCliente2["width"] = 30
    CidadeCliente2.grid(row=7, column=5)

    EstadoCliente2Label = tk.Label(JanelaCliente3, text="Estado do CLiente")
    EstadoCliente2Label.grid(row=8, column=4)

    listEstadoCliente2 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente2
    EstadoCliente2 = ttk.Combobox(JanelaCliente3, values=listEstadoCliente2)
    EstadoCliente2.bind("<Up>", lambda e: CidadeCliente2.focus())
    EstadoCliente2.bind("<Down>", lambda e: CEPCliente2.focus())
    EstadoCliente2.bind("<Return>", lambda e: CEPCliente2.focus())
    EstadoCliente2.set("SP")
    EstadoCliente2["width"] = 27
    EstadoCliente2.grid(row=8, column=5)

    CEPCliente2Label = tk.Label(JanelaCliente3, text="CEP do Cliente")
    CEPCliente2Label.grid(row=9, column=4)

    global CEPCliente2
    CEPCliente2 = tk.Entry(JanelaCliente3)
    CEPCliente2.bind("<Up>", lambda e: EstadoCliente2.focus())
    CEPCliente2.bind("<Down>", lambda e: NacionalidadeCliente2.focus())
    CEPCliente2.bind("<Return>", lambda e: NacionalidadeCliente2.focus())
    CEPCliente2.bind("<KeyRelease>", format_CEP2)
    CEPCliente2["width"] = 30
    CEPCliente2.grid(row=9, column=5)

    NacionalidadeCliente2Label = tk.Label(JanelaCliente3, text="Nacionalidade do Cliente")
    NacionalidadeCliente2Label.grid(row=10, column=4)

    global NacionalidadeCliente2
    NacionalidadeCliente2 = tk.Entry(JanelaCliente3)
    NacionalidadeCliente2.bind("<Up>", lambda e: CEPCliente2.focus())
    NacionalidadeCliente2.bind("<Down>", lambda e: cb_EstCivilCliente2.focus())
    NacionalidadeCliente2.bind("<Return>", lambda e: cb_EstCivilCliente2.focus())
    NacionalidadeCliente2["width"] = 30
    NacionalidadeCliente2.grid(row=10, column=5)

    listEstCivilCliente2 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente2Label = tk.Label(JanelaCliente3, text="Estado Civil Cliente")
    EstCivilCliente2Label.grid(row=11, column=4)

    global cb_EstCivilCliente2
    cb_EstCivilCliente2 = ttk.Combobox(JanelaCliente3, values=listEstCivilCliente2)
    cb_EstCivilCliente2.set("Solteiro(a)")
    cb_EstCivilCliente2.bind("<Up>", lambda e: NacionalidadeCliente2.focus())
    cb_EstCivilCliente2.bind("<Down>", lambda e: ProfissaoCliente2.focus())
    cb_EstCivilCliente2.bind("<Return>", lambda e: ProfissaoCliente2.focus())
    cb_EstCivilCliente2["width"] = 27
    cb_EstCivilCliente2.grid(row=11, column=5)

    ProfissaoCliente2Label = tk.Label(JanelaCliente3, text="Profissão do Cliente")
    ProfissaoCliente2Label.grid(row=12, column=4)

    global ProfissaoCliente2
    ProfissaoCliente2 = tk.Entry(JanelaCliente3)
    ProfissaoCliente2.bind("<Up>", lambda e: cb_EstCivilCliente2.focus())
    ProfissaoCliente2.bind("<Down>", lambda e: EmailCliente2.focus())
    ProfissaoCliente2.bind("<Return>", lambda e: EmailCliente2.focus())
    ProfissaoCliente2["width"] = 30
    ProfissaoCliente2.grid(row=12, column=5)

    EmailCliente2Label = tk.Label(JanelaCliente3, text="Email do Cliente")
    EmailCliente2Label.grid(row=13, column=4)

    global EmailCliente2
    EmailCliente2 = tk.Entry(JanelaCliente3)
    EmailCliente2.bind("<Up>", lambda e: ProfissaoCliente2.focus())
    EmailCliente2.bind("<Down>", lambda e: CelCliente2.focus())
    EmailCliente2.bind("<Return>", lambda e: CelCliente2.focus())
    EmailCliente2["width"] = 30
    EmailCliente2.grid(row=13, column=5)

    CelCliente2Label = tk.Label(JanelaCliente3, text="Celular do Cliente")
    CelCliente2Label.grid(row=14, column=4)

    global CelCliente2
    CelCliente2 = tk.Entry(JanelaCliente3)
    CelCliente2.bind("<Up>", lambda e: EmailCliente2.focus())
    CelCliente2.bind("<Down>", lambda e: NomeCliente3.focus())
    CelCliente2.bind("<Return>", lambda e: NomeCliente3.focus())
    CelCliente2.bind("<KeyRelease>", format_celular2)
    CelCliente2["width"] = 30
    CelCliente2.grid(row=14, column=5)

    # --------------------------------------------------CLIENTE 3 ----------------------------------------------------------------------------------

    label_nome3 = tk.Label(JanelaCliente3, text="Dados do Cliente 3",
                           font=("Arial", "11", "bold"))
    label_nome3.grid(row=0, column=7)

    NomeCLiente3Label = tk.Label(JanelaCliente3, text="Nome do Cliente 3")
    NomeCLiente3Label.grid(row=1, column=7)

    global NomeCliente3
    NomeCliente3 = tk.Entry(JanelaCliente3)
    NomeCliente3.bind("<Up>", lambda e: CelCliente2.focus())
    NomeCliente3.bind("<Down>", lambda e: CPFCliente3.focus())
    NomeCliente3.bind("<Return>", lambda e: CPFCliente3.focus())
    NomeCliente3["width"] = 30
    NomeCliente3.grid(row=1, column=8)

    CPFCliente3Label = tk.Label(JanelaCliente3, text="CPF do Cliente 3")
    CPFCliente3Label.grid(row=2, column=7)

    global CPFCliente3
    CPFCliente3 = tk.Entry(JanelaCliente3)
    CPFCliente3.bind("<Up>", lambda e: NomeCliente3.focus())
    CPFCliente3.bind("<Down>", lambda e: RGCliente3.focus())
    CPFCliente3.bind("<Return>", lambda e: RGCliente3.focus())
    CPFCliente3.bind("<KeyRelease>", format_cpf3)
    CPFCliente3["width"] = 30
    CPFCliente3.grid(row=2, column=8)

    RGCliente3Label = tk.Label(JanelaCliente3, text="RG do Cliente")
    RGCliente3Label.grid(row=3, column=7)

    global RGCliente3
    RGCliente3 = tk.Entry(JanelaCliente3)
    RGCliente3.bind("<Up>", lambda e: CPFCliente3.focus())
    RGCliente3.bind("<Down>", lambda e: EndCliente3.focus())
    RGCliente3.bind("<Return>", lambda e: EndCliente3.focus())
    RGCliente3.bind("<KeyRelease>", format_rg3)
    RGCliente3["width"] = 30
    RGCliente3.grid(row=3, column=8)

    EndCliente3Label = tk.Label(JanelaCliente3, text="Endereço do Cliente")
    EndCliente3Label.grid(row=4, column=7)

    global EndCliente3
    EndCliente3 = tk.Entry(JanelaCliente3)
    EndCliente3.bind("<Up>", lambda e: RGCliente3.focus())
    EndCliente3.bind("<Down>", lambda e: NCliente3.focus())
    EndCliente3.bind("<Return>", lambda e: NCliente3.focus())
    EndCliente3["width"] = 30
    EndCliente3.grid(row=4, column=8)

    NCliente3Label = tk.Label(JanelaCliente3, text="Numero do Cliente")
    NCliente3Label.grid(row=5, column=7)

    global NCliente3
    NCliente3 = tk.Entry(JanelaCliente3)
    NCliente3.bind("<Up>", lambda e: EndCliente3.focus())
    NCliente3.bind("<Down>", lambda e: BairroCliente3.focus())
    NCliente3.bind("<Return>", lambda e: BairroCliente3.focus())
    NCliente3["width"] = 30
    NCliente3.grid(row=5, column=8)

    BairroCliente3Label = tk.Label(JanelaCliente3, text="Bairro do Cliente")
    BairroCliente3Label.grid(row=6, column=7)

    global BairroCliente3
    BairroCliente3 = tk.Entry(JanelaCliente3)
    BairroCliente3.bind("<Up>", lambda e: NCliente3.focus())
    BairroCliente3.bind("<Down>", lambda e: CidadeCliente3.focus())
    BairroCliente3.bind("<Return>", lambda e: CidadeCliente3.focus())
    BairroCliente3["width"] = 30
    BairroCliente3.grid(row=6, column=8)

    CidadeCliente3Label = tk.Label(JanelaCliente3, text="Cidade do Cliente")
    CidadeCliente3Label.grid(row=7, column=7)

    global CidadeCliente3
    CidadeCliente3 = tk.Entry(JanelaCliente3)
    CidadeCliente3.bind("<Up>", lambda e: BairroCliente3.focus())
    CidadeCliente3.bind("<Down>", lambda e: EstadoCliente3.focus())
    CidadeCliente3.bind("<Return>", lambda e: EstadoCliente3.focus())
    CidadeCliente3["width"] = 30
    CidadeCliente3.grid(row=7, column=8)

    EstadoCliente3Label = tk.Label(JanelaCliente3, text="Estado do CLiente")
    EstadoCliente3Label.grid(row=8, column=7)

    listEstadoCliente3 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente3
    EstadoCliente3 = ttk.Combobox(JanelaCliente3, values=listEstadoCliente3)
    EstadoCliente3.bind("<Up>", lambda e: CidadeCliente3.focus())
    EstadoCliente3.bind("<Down>", lambda e: CEPCliente3.focus())
    EstadoCliente3.bind("<Return>", lambda e: CEPCliente3.focus())
    EstadoCliente3.set("SP")
    EstadoCliente3["width"] = 27
    EstadoCliente3.grid(row=8, column=8)

    CEPCliente3Label = tk.Label(JanelaCliente3, text="CEP do Cliente")
    CEPCliente3Label.grid(row=9, column=7)

    global CEPCliente3
    CEPCliente3 = tk.Entry(JanelaCliente3)
    CEPCliente3.bind("<Up>", lambda e: EstadoCliente3.focus())
    CEPCliente3.bind("<Down>", lambda e: NacionalidadeCliente3.focus())
    CEPCliente3.bind("<Return>", lambda e: NacionalidadeCliente3.focus())
    CEPCliente3.bind("<KeyRelease>", format_CEP3)
    CEPCliente3["width"] = 30
    CEPCliente3.grid(row=9, column=8)

    NacionalidadeCliente3Label = tk.Label(JanelaCliente3, text="Nacionalidade do Cliente")
    NacionalidadeCliente3Label.grid(row=10, column=7)

    global NacionalidadeCliente3
    NacionalidadeCliente3 = tk.Entry(JanelaCliente3)
    NacionalidadeCliente3.bind("<Up>", lambda e: CEPCliente3.focus())
    NacionalidadeCliente3.bind("<Down>", lambda e: cb_EstCivilCliente3.focus())
    NacionalidadeCliente3.bind("<Return>", lambda e: cb_EstCivilCliente3.focus())
    NacionalidadeCliente3["width"] = 30
    NacionalidadeCliente3.grid(row=10, column=8)

    listEstCivilCliente3 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente3Label = tk.Label(JanelaCliente3, text="Estado Civil Cliente")
    EstCivilCliente3Label.grid(row=11, column=7)

    global cb_EstCivilCliente3
    cb_EstCivilCliente3 = ttk.Combobox(JanelaCliente3, values=listEstCivilCliente3)
    cb_EstCivilCliente3.bind("<Up>", lambda e: NacionalidadeCliente3.focus())
    cb_EstCivilCliente3.bind("<Down>", lambda e: ProfissaoCliente3.focus())
    cb_EstCivilCliente3.bind("<Return>", lambda e: ProfissaoCliente3.focus())
    cb_EstCivilCliente3.set("Solteiro(a)")
    cb_EstCivilCliente3["width"] = 27
    cb_EstCivilCliente3.grid(row=11, column=8)

    ProfissaoCliente3Label = tk.Label(JanelaCliente3, text="Profissão do Cliente")
    ProfissaoCliente3Label.grid(row=12, column=7)

    global ProfissaoCliente3
    ProfissaoCliente3 = tk.Entry(JanelaCliente3)
    ProfissaoCliente3.bind("<Up>", lambda e: cb_EstCivilCliente3.focus())
    ProfissaoCliente3.bind("<Down>", lambda e: EmailCliente3.focus())
    ProfissaoCliente3.bind("<Return>", lambda e: EmailCliente3.focus())
    ProfissaoCliente3["width"] = 30
    ProfissaoCliente3.grid(row=12, column=8)

    EmailCliente3Label = tk.Label(JanelaCliente3, text="Email do Cliente")
    EmailCliente3Label.grid(row=13, column=7)

    global EmailCliente3
    EmailCliente3 = tk.Entry(JanelaCliente3)
    EmailCliente3.bind("<Up>", lambda e: ProfissaoCliente3.focus())
    EmailCliente3.bind("<Down>", lambda e: CelCliente3.focus())
    EmailCliente3.bind("<Return>", lambda e: CelCliente3.focus())
    EmailCliente3["width"] = 30
    EmailCliente3.grid(row=13, column=8)

    CelCliente3Label = tk.Label(JanelaCliente3, text="Celular do Cliente")
    CelCliente3Label.grid(row=14, column=7)

    global CelCliente3
    CelCliente3 = tk.Entry(JanelaCliente3)
    CelCliente3.bind("<Up>", lambda e: EmailCliente3.focus())
    CelCliente3.bind("<Down>", lambda e: NomeCliente1.focus())
    CelCliente3.bind("<Return>", lambda e: confirmar3)
    CelCliente3.bind("<KeyRelease>", format_celular3)
    CelCliente3["width"] = 30
    CelCliente3.grid(row=14, column=8)

    bntVoltar3 = tk.Button(JanelaCliente3, text="Voltar", command=volta3)
    bntVoltar3.grid(row=16, column=2)

    bntCliente3 = tk.Button(JanelaCliente3, text="Confirmar")
    bntCliente3["command"] = confirmar3
    bntCliente3.grid(row=16, column=4)

    bntCopiar = tk.Button(JanelaCliente3, text="Copiar Cliente 1 para 2")
    bntCopiar["command"] = copiar21
    bntCopiar.grid(row=16, column=5)

    bntCopiar2 = tk.Button(JanelaCliente3, text="Copiar Cliente 2 para 3")
    bntCopiar2["command"] = copiar32
    bntCopiar2.grid(row=16, column=8)


# Janela para quatro clientes
def clientefour():
    global JanelaCliente4
    JanelaCliente4 = tk.Toplevel()
    JanelaCliente4.title('Cadastro Cliente')
    JanelaCliente4.iconbitmap('image/logo.ico')
    JanelaCliente4.configure(bg='white')

    janela2.withdraw()

    JanelaCliente4.protocol("WM_DELETE_WINDOW", on_closing)

    label_nome = tk.Label(JanelaCliente4, text="Dados do CLiente 1",
                          font=("Arial", "11", "bold"))
    label_nome.grid(row=0, column=1)

    NomeCLiente1Label = tk.Label(JanelaCliente4, text="Nome do Cliente ")
    NomeCLiente1Label.grid(row=1, column=1)

    global NomeCliente1
    NomeCliente1 = tk.Entry(JanelaCliente4)
    NomeCliente1.focus()
    NomeCliente1.bind("<Up>", lambda e: NCliente1.focus())
    NomeCliente1.bind("<Down>", lambda e: CPFCliente1.focus())
    NomeCliente1.bind("<Return>", lambda e: CPFCliente1.focus())
    NomeCliente1["width"] = 30
    NomeCliente1.grid(row=1, column=2)

    CPFCliente1Label = tk.Label(JanelaCliente4, text="CPF do Cliente")
    CPFCliente1Label.grid(row=2, column=1)

    global CPFCliente1
    CPFCliente1 = tk.Entry(JanelaCliente4)
    CPFCliente1.bind("<KeyPress>", lambda e: RGCliente1.focus() if e.char == '\r' else None)
    CPFCliente1.bind("<KeyRelease>", format_cpf1)
    CPFCliente1["width"] = 30
    CPFCliente1.grid(row=2, column=2)

    RGCliente1Label = tk.Label(JanelaCliente4, text="RG do Cliente")
    RGCliente1Label.grid(row=3, column=1)

    global RGCliente1
    RGCliente1 = tk.Entry(JanelaCliente4)
    RGCliente1.bind("<KeyPress>", lambda e: EndCliente1.focus() if e.char == '\r' else None)
    RGCliente1.bind("<KeyRelease>", format_rg1)
    RGCliente1["width"] = 30
    RGCliente1.grid(row=3, column=2)

    EndCliente1Label = tk.Label(JanelaCliente4, text="Endereço do Cliente")
    EndCliente1Label.grid(row=4, column=1)

    global EndCliente1
    EndCliente1 = tk.Entry(JanelaCliente4)
    EndCliente1.bind("<KeyPress>", lambda e: NCliente1.focus() if e.char == '\r' else None)
    EndCliente1["width"] = 30
    EndCliente1.grid(row=4, column=2)

    NCliente1Label = tk.Label(JanelaCliente4, text="Numero do Cliente")
    NCliente1Label.grid(row=5, column=1)

    global NCliente1
    NCliente1 = tk.Entry(JanelaCliente4)
    NCliente1.bind("<KeyPress>", lambda e: BairroCliente1.focus() if e.char == '\r' else None)
    NCliente1["width"] = 30
    NCliente1.grid(row=5, column=2)

    BairroCliente1Label = tk.Label(JanelaCliente4, text="Bairro do Cliente")
    BairroCliente1Label.grid(row=6, column=1)

    global BairroCliente1
    BairroCliente1 = tk.Entry(JanelaCliente4)
    BairroCliente1.bind("<KeyPress>", lambda e: CidadeCliente1.focus() if e.char == '\r' else None)
    BairroCliente1["width"] = 30
    BairroCliente1.grid(row=6, column=2)

    CidadeCliente1Label = tk.Label(JanelaCliente4, text="Cidade do Cliente")
    CidadeCliente1Label.grid(row=7, column=1)

    global CidadeCliente1
    CidadeCliente1 = tk.Entry(JanelaCliente4)
    CidadeCliente1.bind("<KeyPress>", lambda e: EstadoCliente1.focus() if e.char == '\r' else None)
    CidadeCliente1["width"] = 30
    CidadeCliente1.grid(row=7, column=2)

    EstadoCliente1Label = tk.Label(JanelaCliente4, text="Estado do CLiente")
    EstadoCliente1Label.grid(row=8, column=1)

    listEstadoCliente1 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente1
    EstadoCliente1 = ttk.Combobox(JanelaCliente4, values=listEstadoCliente1)
    EstadoCliente1.bind("<KeyPress>", lambda e: CEPCliente1.focus() if e.char == '\r' else None)
    EstadoCliente1.set("SP")
    EstadoCliente1["width"] = 27
    EstadoCliente1.grid(row=8, column=2)

    CEPCliente1Label = tk.Label(JanelaCliente4, text="CEP do Cliente")
    CEPCliente1Label.grid(row=9, column=1)

    global CEPCliente1
    CEPCliente1 = tk.Entry(JanelaCliente4)
    CEPCliente1.bind("<KeyPress>", lambda e: NacionalidadeCliente1.focus() if e.char == '\r' else None)
    CEPCliente1.bind("<KeyRelease>", format_CEP1)
    CEPCliente1["width"] = 30
    CEPCliente1.grid(row=9, column=2)

    NacionalidadeCliente1Label = tk.Label(JanelaCliente4, text="Nacionalidade do Cliente")
    NacionalidadeCliente1Label.grid(row=10, column=1)

    global NacionalidadeCliente1
    NacionalidadeCliente1 = tk.Entry(JanelaCliente4)
    NacionalidadeCliente1.bind("<KeyPress>", lambda e: cb_EstCivilCliente1.focus() if e.char == '\r' else None)
    NacionalidadeCliente1["width"] = 30
    NacionalidadeCliente1.grid(row=10, column=2)

    listEstCivilCliente1 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente1Label = tk.Label(JanelaCliente4, text="Estado Civil Cliente")
    EstCivilCliente1Label.grid(row=11, column=1)

    global cb_EstCivilCliente1
    cb_EstCivilCliente1 = ttk.Combobox(JanelaCliente4, values=listEstCivilCliente1)
    cb_EstCivilCliente1.bind("<KeyPress>", lambda e: ProfissaoCliente1.focus() if e.char == '\r' else None)
    cb_EstCivilCliente1.set("Solteiro(a)")
    cb_EstCivilCliente1["width"] = 27
    cb_EstCivilCliente1.grid(row=11, column=2)

    ProfissaoCliente1Label = tk.Label(JanelaCliente4, text="Profissão do Cliente")
    ProfissaoCliente1Label.grid(row=12, column=1)

    global ProfissaoCliente1
    ProfissaoCliente1 = tk.Entry(JanelaCliente4)
    ProfissaoCliente1.bind("<KeyPress>", lambda e: EmailCliente1.focus() if e.char == '\r' else None)
    ProfissaoCliente1["width"] = 30
    ProfissaoCliente1.grid(row=12, column=2)

    EmailCliente1Label = tk.Label(JanelaCliente4, text="Email do Cliente")
    EmailCliente1Label.grid(row=13, column=1)

    global EmailCliente1
    EmailCliente1 = tk.Entry(JanelaCliente4)
    EmailCliente1.bind("<KeyPress>", lambda e: CelCliente1.focus() if e.char == '\r' else None)
    EmailCliente1["width"] = 30
    EmailCliente1.grid(row=13, column=2)

    CelCliente1Label = tk.Label(JanelaCliente4, text="Celular do Cliente")
    CelCliente1Label.grid(row=14, column=1)

    global CelCliente1
    CelCliente1 = tk.Entry(JanelaCliente4)
    CelCliente1.bind("<KeyPress>", lambda e: NomeCliente2.focus() if e.char == '\r' else None)
    CelCliente1.bind("<KeyRelease>", format_celular1)
    CelCliente1["width"] = 30
    CelCliente1.grid(row=14, column=2)

    # --------------------------------------------------CLIENTE 2 ----------------------------------------------------------------------------------

    label_nome2 = tk.Label(JanelaCliente4, text="Dados do Cliente 2",
                           font=("Arial", "11", "bold"))
    label_nome2.grid(row=0, column=4)

    NomeCLiente2Label = tk.Label(JanelaCliente4, text="Nome do Cliente 2")
    NomeCLiente2Label.grid(row=1, column=4)

    global NomeCliente2
    NomeCliente2 = tk.Entry(JanelaCliente4)
    NomeCliente2.bind("<KeyPress>", lambda e: CPFCliente2.focus() if e.char == '\r' else None)
    NomeCliente2["width"] = 30
    NomeCliente2.grid(row=1, column=5)

    CPFCliente2Label = tk.Label(JanelaCliente4, text="CPF do Cliente 2")
    CPFCliente2Label.grid(row=2, column=4)

    global CPFCliente2
    CPFCliente2 = tk.Entry(JanelaCliente4)
    CPFCliente2.bind("<KeyPress>", lambda e: RGCliente2.focus() if e.char == '\r' else None)
    CPFCliente2.bind("<KeyRelease>", format_cpf2)
    CPFCliente2["width"] = 30
    CPFCliente2.grid(row=2, column=5)

    RGCliente2Label = tk.Label(JanelaCliente4, text="RG do Cliente")
    RGCliente2Label.grid(row=3, column=4)

    global RGCliente2
    RGCliente2 = tk.Entry(JanelaCliente4)
    RGCliente2.bind("<KeyPress>", lambda e: EndCliente2.focus() if e.char == '\r' else None)
    RGCliente2.bind("<KeyRelease>", format_rg2)
    RGCliente2["width"] = 30
    RGCliente2.grid(row=3, column=5)

    EndCliente2Label = tk.Label(JanelaCliente4, text="Endereço do Cliente")
    EndCliente2Label.grid(row=4, column=4)

    global EndCliente2
    EndCliente2 = tk.Entry(JanelaCliente4)
    EndCliente2.bind("<KeyPress>", lambda e: NCliente2.focus() if e.char == '\r' else None)
    EndCliente2["width"] = 30
    EndCliente2.grid(row=4, column=5)

    NCliente2Label = tk.Label(JanelaCliente4, text="Numero do Cliente")
    NCliente2Label.grid(row=5, column=4)

    global NCliente2
    NCliente2 = tk.Entry(JanelaCliente4)
    NCliente2.bind("<KeyPress>", lambda e: BairroCliente2.focus() if e.char == '\r' else None)
    NCliente2["width"] = 30
    NCliente2.grid(row=5, column=5)

    BairroCliente2Label = tk.Label(JanelaCliente4, text="Bairro do Cliente")
    BairroCliente2Label.grid(row=6, column=4)

    global BairroCliente2
    BairroCliente2 = tk.Entry(JanelaCliente4)
    BairroCliente2.bind("<KeyPress>", lambda e: CidadeCliente2.focus() if e.char == '\r' else None)
    BairroCliente2["width"] = 30
    BairroCliente2.grid(row=6, column=5)

    CidadeCliente2Label = tk.Label(JanelaCliente4, text="Cidade do Cliente")
    CidadeCliente2Label.grid(row=7, column=4)

    global CidadeCliente2
    CidadeCliente2 = tk.Entry(JanelaCliente4)
    CidadeCliente2.bind("<KeyPress>", lambda e: EstadoCliente2.focus() if e.char == '\r' else None)
    CidadeCliente2["width"] = 30
    CidadeCliente2.grid(row=7, column=5)

    EstadoCliente2Label = tk.Label(JanelaCliente4, text="Estado do CLiente")
    EstadoCliente2Label.grid(row=8, column=4)

    listEstadoCliente2 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente2
    EstadoCliente2 = ttk.Combobox(JanelaCliente4, values=listEstadoCliente2)
    EstadoCliente2.bind("<KeyPress>", lambda e: CEPCliente2.focus() if e.char == '\r' else None)
    EstadoCliente2.set("SP")
    EstadoCliente2["width"] = 27
    EstadoCliente2.grid(row=8, column=5)

    CEPCliente2Label = tk.Label(JanelaCliente4, text="CEP do Cliente")
    CEPCliente2Label.grid(row=9, column=4)

    global CEPCliente2
    CEPCliente2 = tk.Entry(JanelaCliente4)
    CEPCliente2.bind("<KeyPress>", lambda e: NacionalidadeCliente2.focus() if e.char == '\r' else None)
    CEPCliente2.bind("<KeyRelease>", format_CEP2)
    CEPCliente2["width"] = 30
    CEPCliente2.grid(row=9, column=5)

    NacionalidadeCliente2Label = tk.Label(JanelaCliente4, text="Nacionalidade do Cliente")
    NacionalidadeCliente2Label.grid(row=10, column=4)

    global NacionalidadeCliente2
    NacionalidadeCliente2 = tk.Entry(JanelaCliente4)
    NacionalidadeCliente2.bind("<KeyPress>", lambda e: cb_EstCivilCliente2.focus() if e.char == '\r' else None)
    NacionalidadeCliente2["width"] = 30
    NacionalidadeCliente2.grid(row=10, column=5)

    listEstCivilCliente2 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente2Label = tk.Label(JanelaCliente4, text="Estado Civil Cliente")
    EstCivilCliente2Label.grid(row=11, column=4)

    global cb_EstCivilCliente2
    cb_EstCivilCliente2 = ttk.Combobox(JanelaCliente4, values=listEstCivilCliente2)
    cb_EstCivilCliente2.set("Solteiro(a)")
    cb_EstCivilCliente2.bind("<KeyPress>", lambda e: ProfissaoCliente2.focus() if e.char == '\r' else None)
    cb_EstCivilCliente2["width"] = 27
    cb_EstCivilCliente2.grid(row=11, column=5)

    ProfissaoCliente2Label = tk.Label(JanelaCliente4, text="Profissão do Cliente")
    ProfissaoCliente2Label.grid(row=12, column=4)

    global ProfissaoCliente2
    ProfissaoCliente2 = tk.Entry(JanelaCliente4)
    ProfissaoCliente2.bind("<KeyPress>", lambda e: EmailCliente2.focus() if e.char == '\r' else None)
    ProfissaoCliente2["width"] = 30
    ProfissaoCliente2.grid(row=12, column=5)

    EmailCliente2Label = tk.Label(JanelaCliente4, text="Email do Cliente")
    EmailCliente2Label.grid(row=13, column=4)

    global EmailCliente2
    EmailCliente2 = tk.Entry(JanelaCliente4)
    EmailCliente2.bind("<KeyPress>", lambda e: CelCliente2.focus() if e.char == '\r' else None)
    EmailCliente2["width"] = 30
    EmailCliente2.grid(row=13, column=5)

    CelCliente2Label = tk.Label(JanelaCliente4, text="Celular do Cliente")
    CelCliente2Label.grid(row=14, column=4)

    global CelCliente2
    CelCliente2 = tk.Entry(JanelaCliente4)
    CelCliente2.bind("<KeyPress>", lambda e: NomeCliente3.focus() if e.char == '\r' else None)
    CelCliente2.bind("<KeyRelease>", format_celular2)
    CelCliente2["width"] = 30
    CelCliente2.grid(row=14, column=5)

    # --------------------------------------------------CLIENTE 3 ----------------------------------------------------------------------------------

    label_nome3 = tk.Label(JanelaCliente4, text="Dados do Cliente 3",
                           font=("Arial", "11", "bold"))
    label_nome3.grid(row=0, column=7)

    NomeCLiente3Label = tk.Label(JanelaCliente4, text="Nome do Cliente 3")
    NomeCLiente3Label.grid(row=1, column=7)

    global NomeCliente3
    NomeCliente3 = tk.Entry(JanelaCliente4)
    NomeCliente3.bind("<KeyPress>", lambda e: CPFCliente3.focus() if e.char == '\r' else None)
    NomeCliente3["width"] = 30
    NomeCliente3.grid(row=1, column=8)

    CPFCliente3Label = tk.Label(JanelaCliente4, text="CPF do Cliente 3")
    CPFCliente3Label.grid(row=2, column=7)

    global CPFCliente3
    CPFCliente3 = tk.Entry(JanelaCliente4)
    CPFCliente3.bind("<KeyPress>", lambda e: RGCliente3.focus() if e.char == '\r' else None)
    CPFCliente3.bind("<KeyRelease>", format_celular3)
    CPFCliente3["width"] = 30
    CPFCliente3.grid(row=2, column=8)

    RGCliente3Label = tk.Label(JanelaCliente4, text="RG do Cliente")
    RGCliente3Label.grid(row=3, column=7)

    global RGCliente3
    RGCliente3 = tk.Entry(JanelaCliente4)
    RGCliente3.bind("<KeyPress>", lambda e: EndCliente3.focus() if e.char == '\r' else None)
    RGCliente3.bind("<KeyRelease>", format_rg3)
    RGCliente3["width"] = 30
    RGCliente3.grid(row=3, column=8)

    EndCliente3Label = tk.Label(JanelaCliente4, text="Endereço do Cliente")
    EndCliente3Label.grid(row=4, column=7)

    global EndCliente3
    EndCliente3 = tk.Entry(JanelaCliente4)
    EndCliente3.bind("<KeyPress>", lambda e: NCliente3.focus() if e.char == '\r' else None)
    EndCliente3["width"] = 30
    EndCliente3.grid(row=4, column=8)

    NCliente3Label = tk.Label(JanelaCliente4, text="Numero do Cliente")
    NCliente3Label.grid(row=5, column=7)

    global NCliente3
    NCliente3 = tk.Entry(JanelaCliente4)
    NCliente3.bind("<KeyPress>", lambda e: BairroCliente3.focus() if e.char == '\r' else None)
    NCliente3["width"] = 30
    NCliente3.grid(row=5, column=8)

    BairroCliente3Label = tk.Label(JanelaCliente4, text="Bairro do Cliente")
    BairroCliente3Label.grid(row=6, column=7)

    global BairroCliente3
    BairroCliente3 = tk.Entry(JanelaCliente4)
    BairroCliente3.bind("<KeyPress>", lambda e: CidadeCliente3.focus() if e.char == '\r' else None)
    BairroCliente3["width"] = 30
    BairroCliente3.grid(row=6, column=8)

    CidadeCliente3Label = tk.Label(JanelaCliente4, text="Cidade do Cliente")
    CidadeCliente3Label.grid(row=7, column=7)

    global CidadeCliente3
    CidadeCliente3 = tk.Entry(JanelaCliente4)
    CidadeCliente3.bind("<KeyPress>", lambda e: EstadoCliente3.focus() if e.char == '\r' else None)
    CidadeCliente3["width"] = 30
    CidadeCliente3.grid(row=7, column=8)

    EstadoCliente3Label = tk.Label(JanelaCliente4, text="Estado do CLiente")
    EstadoCliente3Label.grid(row=8, column=7)

    listEstadoCliente3 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente3
    EstadoCliente3 = ttk.Combobox(JanelaCliente4, values=listEstadoCliente3)
    EstadoCliente3.bind("<KeyPress>", lambda e: CEPCliente3.focus() if e.char == '\r' else None)
    EstadoCliente3.set("SP")
    EstadoCliente3["width"] = 27
    EstadoCliente3.grid(row=8, column=8)

    CEPCliente3Label = tk.Label(JanelaCliente4, text="CEP do Cliente")
    CEPCliente3Label.grid(row=9, column=7)

    global CEPCliente3
    CEPCliente3 = tk.Entry(JanelaCliente4)
    CEPCliente3.bind("<KeyPress>", lambda e: NacionalidadeCliente3.focus() if e.char == '\r' else None)
    CEPCliente3.bind("<KeyRelease>", format_CEP3)
    CEPCliente3["width"] = 30
    CEPCliente3.grid(row=9, column=8)

    NacionalidadeCliente3Label = tk.Label(JanelaCliente4, text="Nacionalidade do Cliente")
    NacionalidadeCliente3Label.grid(row=10, column=7)

    global NacionalidadeCliente3
    NacionalidadeCliente3 = tk.Entry(JanelaCliente4)
    NacionalidadeCliente3.bind("<KeyPress>", lambda e: cb_EstCivilCliente3.focus() if e.char == '\r' else None)
    NacionalidadeCliente3["width"] = 30
    NacionalidadeCliente3.grid(row=10, column=8)

    listEstCivilCliente3 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente3Label = tk.Label(JanelaCliente4, text="Estado Civil Cliente")
    EstCivilCliente3Label.grid(row=11, column=7)

    global cb_EstCivilCliente3
    cb_EstCivilCliente3 = ttk.Combobox(JanelaCliente4, values=listEstCivilCliente3)
    cb_EstCivilCliente3.bind("<KeyPress>", lambda e: ProfissaoCliente3.focus() if e.char == '\r' else None)
    cb_EstCivilCliente3.set("Solteiro(a)")
    cb_EstCivilCliente3["width"] = 27
    cb_EstCivilCliente3.grid(row=11, column=8)

    ProfissaoCliente3Label = tk.Label(JanelaCliente4, text="Profissão do Cliente")
    ProfissaoCliente3Label.grid(row=12, column=7)

    global ProfissaoCliente3
    ProfissaoCliente3 = tk.Entry(JanelaCliente4)
    ProfissaoCliente3.bind("<KeyPress>", lambda e: EmailCliente3.focus() if e.char == '\r' else None)
    ProfissaoCliente3["width"] = 30
    ProfissaoCliente3.grid(row=12, column=8)

    EmailCliente3Label = tk.Label(JanelaCliente4, text="Email do Cliente")
    EmailCliente3Label.grid(row=13, column=7)

    global EmailCliente3
    EmailCliente3 = tk.Entry(JanelaCliente4)
    EmailCliente3.bind("<KeyPress>", lambda e: CelCliente3.focus() if e.char == '\r' else None)
    EmailCliente3["width"] = 30
    EmailCliente3.grid(row=13, column=8)

    CelCliente3Label = tk.Label(JanelaCliente4, text="Celular do Cliente")
    CelCliente3Label.grid(row=14, column=7)

    global CelCliente3
    CelCliente3 = tk.Entry(JanelaCliente4)
    CelCliente3.bind("<KeyPress>", lambda e: NomeCliente4.focus() if e.char == '\r' else None)
    CelCliente3.bind("<KeyRelease>", format_celular3)
    CelCliente3["width"] = 30
    CelCliente3.grid(row=14, column=8)

    # --------------------------------------------------CLIENTE 4 ----------------------------------------------------------------------------------

    label_nome4 = tk.Label(JanelaCliente4, text="Dados do Cliente 4",
                           font=("Arial", "11", "bold"))
    label_nome4.grid(row=0, column=10)

    NomeCLiente4Label = tk.Label(JanelaCliente4, text="Nome do Cliente 4")
    NomeCLiente4Label.grid(row=1, column=10)

    global NomeCliente4
    NomeCliente4 = tk.Entry(JanelaCliente4)
    NomeCliente4.bind("<KeyPress>", lambda e: CPFCliente4.focus() if e.char == '\r' else None)
    NomeCliente4["width"] = 30
    NomeCliente4.grid(row=1, column=11)

    CPFCliente4Label = tk.Label(JanelaCliente4, text="CPF do Cliente 4")
    CPFCliente4Label.grid(row=2, column=10)

    global CPFCliente4
    CPFCliente4 = tk.Entry(JanelaCliente4)
    CPFCliente4.bind("<KeyPress>", lambda e: RGCliente4.focus() if e.char == '\r' else None)
    CPFCliente4.bind("<KeyRelease>", format_cpf4)
    CPFCliente4["width"] = 30
    CPFCliente4.grid(row=2, column=11)

    RGCliente4Label = tk.Label(JanelaCliente4, text="RG do Cliente")
    RGCliente4Label.grid(row=3, column=10)

    global RGCliente4
    RGCliente4 = tk.Entry(JanelaCliente4)
    RGCliente4.bind("<KeyPress>", lambda e: EndCliente4.focus() if e.char == '\r' else None)
    RGCliente4.bind("<KeyRelease>", format_rg4)
    RGCliente4["width"] = 30
    RGCliente4.grid(row=3, column=11)

    EndCliente4Label = tk.Label(JanelaCliente4, text="Endereço do Cliente")
    EndCliente4Label.grid(row=4, column=10)

    global EndCliente4
    EndCliente4 = tk.Entry(JanelaCliente4)
    EndCliente4.bind("<KeyPress>", lambda e: NCliente4.focus() if e.char == '\r' else None)
    EndCliente4["width"] = 30
    EndCliente4.grid(row=4, column=11)

    NCliente4Label = tk.Label(JanelaCliente4, text="Numero do Cliente")
    NCliente4Label.grid(row=5, column=10)

    global NCliente4
    NCliente4 = tk.Entry(JanelaCliente4)
    NCliente4.bind("<KeyPress>", lambda e: BairroCliente4.focus() if e.char == '\r' else None)
    NCliente4["width"] = 30
    NCliente4.grid(row=5, column=11)

    BairroCliente4Label = tk.Label(JanelaCliente4, text="Bairro do Cliente")
    BairroCliente4Label.grid(row=6, column=10)

    global BairroCliente4
    BairroCliente4 = tk.Entry(JanelaCliente4)
    BairroCliente4.bind("<KeyPress>", lambda e: CidadeCliente4.focus() if e.char == '\r' else None)
    BairroCliente4["width"] = 30
    BairroCliente4.grid(row=6, column=11)

    CidadeCliente4Label = tk.Label(JanelaCliente4, text="Cidade do Cliente")
    CidadeCliente4Label.grid(row=7, column=10)

    global CidadeCliente4
    CidadeCliente4 = tk.Entry(JanelaCliente4)
    CidadeCliente4.bind("<KeyPress>", lambda e: EstadoCliente4.focus() if e.char == '\r' else None)
    CidadeCliente4["width"] = 30
    CidadeCliente4.grid(row=7, column=11)

    EstadoCliente4Label = tk.Label(JanelaCliente4, text="Estado do CLiente")
    EstadoCliente4Label.grid(row=8, column=10)

    listEstadoCliente4 = ["AC", "AL", "AP", "AM", "BA", "CE", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR",
                          "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO", "DF"]

    global EstadoCliente4
    EstadoCliente4 = ttk.Combobox(JanelaCliente4, values=listEstadoCliente4)
    EstadoCliente4.bind("<KeyPress>", lambda e: CEPCliente4.focus() if e.char == '\r' else None)
    EstadoCliente4.set("SP")
    EstadoCliente4["width"] = 27
    EstadoCliente4.grid(row=8, column=11)

    CEPCliente4Label = tk.Label(JanelaCliente4, text="CEP do Cliente")
    CEPCliente4Label.grid(row=9, column=10)

    global CEPCliente4
    CEPCliente4 = tk.Entry(JanelaCliente4)
    CEPCliente4.bind("<KeyPress>", lambda e: NacionalidadeCliente4.focus() if e.char == '\r' else None)
    CEPCliente4.bind("<KeyRelease>", format_CEP4)
    CEPCliente4["width"] = 30
    CEPCliente4.grid(row=9, column=11)

    NacionalidadeCliente4Label = tk.Label(JanelaCliente4, text="Nacionalidade do Cliente")
    NacionalidadeCliente4Label.grid(row=10, column=10)

    global NacionalidadeCliente4
    NacionalidadeCliente4 = tk.Entry(JanelaCliente4)
    NacionalidadeCliente4.bind("<KeyPress>", lambda e: cb_EstCivilCliente4.focus() if e.char == '\r' else None)
    NacionalidadeCliente4["width"] = 30
    NacionalidadeCliente4.grid(row=10, column=11)

    listEstCivilCliente4 = ["Solteiro(a)", "Casado(a)", "Viúvo(a)", "Divorciado(a)"]

    EstCivilCliente4Label = tk.Label(JanelaCliente4, text="Estado Civil Cliente")
    EstCivilCliente4Label.grid(row=11, column=10)

    global cb_EstCivilCliente4
    cb_EstCivilCliente4 = ttk.Combobox(JanelaCliente4, values=listEstCivilCliente4)
    cb_EstCivilCliente4.bind("<KeyPress>", lambda e: ProfissaoCliente4.focus() if e.char == '\r' else None)
    cb_EstCivilCliente4.set("Solteiro(a)")
    cb_EstCivilCliente4["width"] = 27
    cb_EstCivilCliente4.grid(row=11, column=11)

    ProfissaoCliente4Label = tk.Label(JanelaCliente4, text="Profissão do Cliente")
    ProfissaoCliente4Label.grid(row=12, column=10)

    global ProfissaoCliente4
    ProfissaoCliente4 = tk.Entry(JanelaCliente4)
    ProfissaoCliente4.bind("<KeyPress>", lambda e: EmailCliente4.focus() if e.char == '\r' else None)
    ProfissaoCliente4["width"] = 30
    ProfissaoCliente4.grid(row=12, column=11)

    EmailCliente4Label = tk.Label(JanelaCliente4, text="Email do Cliente")
    EmailCliente4Label.grid(row=13, column=10)

    global EmailCliente4
    EmailCliente4 = tk.Entry(JanelaCliente4)
    EmailCliente4.bind("<KeyPress>", lambda e: CelCliente4.focus() if e.char == '\r' else None)
    EmailCliente4["width"] = 30
    EmailCliente4.grid(row=13, column=11)

    CelCliente4Label = tk.Label(JanelaCliente4, text="Celular do Cliente")
    CelCliente4Label.grid(row=14, column=10)

    global CelCliente4
    CelCliente4 = tk.Entry(JanelaCliente4)
    CelCliente4.bind("<KeyPress>", lambda e: confirmar4() if e.char == '\r' else None)
    CelCliente4.bind("<KeyRelease>", format_celular4)
    CelCliente4["width"] = 30
    CelCliente4.grid(row=14, column=11)

    bntVoltar4 = tk.Button(JanelaCliente4, text="Voltar", command=volta4)
    bntVoltar4.grid(row=16, column=2)

    bntCliente4 = tk.Button(JanelaCliente4, text="Confirmar")
    bntCliente4["command"] = confirmar4
    bntCliente4.grid(row=16, column=4)

    bntCopiar = tk.Button(JanelaCliente4, text="Copiar Cliente 1 para 2")
    bntCopiar["command"] = copiar21
    bntCopiar.grid(row=16, column=5)

    bntCopiar2 = tk.Button(JanelaCliente4, text="Copiar Cliente 2 para 3")
    bntCopiar2["command"] = copiar32
    bntCopiar2.grid(row=16, column=8)

    bntCopiar3 = tk.Button(JanelaCliente4, text="Copiar Cliente 3 para 4")
    bntCopiar3["command"] = copiar43
    bntCopiar3.grid(row=16, column=11)


# Janela para fazer doc de um cliente
def confirmar1():
    # variaveis da obra
    var_EndObra = endobra.get()
    var_NObra = NumObra.get()
    var_BairroObra = BairroObra.get()
    var_LoteObra = LoteObra.get()
    var_QuadraObra = QuadraObra.get()
    var_CidadeObra = CidadeObra.get()
    var_QuarteiraoObra = QuarteiraoObra.get()
    var_TipoObra = TipoObra.get()
    var_AreaObra = AreaObra.get()
    var_ArtObra = ArtObra.get()
    var_ValorContrato = ValorContrato.get()
    var_ValorContratoextenso = ValorContratoextenso.get()
    var_ParcelContrato = ParcelContrato.get()
    var_ParcelContratoextenso = ParcelContratoextenso.get()
    var_ValorParcelaContrato = ValorParcelaContrato.get()
    var_ValorParcelaContratoextenso = ValorParcelaContratoextenso.get()
    var_DiaVencimento = DiaVencimento.get()
    var_InicioContrato = InicioContrato.get()
    var_Visita = Visita.get()
    var_Visitaextenso = Visitaextenso.get()

    # variaveis do cliente
    var_NomeCliente1 = NomeCliente1.get()
    var_CPFCliente1 = CPFCliente1.get()
    var_RGCliente1 = RGCliente1.get()
    var_EndCliente1 = EndCliente1.get()
    var_NCliente1 = NCliente1.get()
    var_BairroCliente1 = BairroCliente1.get()
    var_CidadeCliente1 = CidadeCliente1.get()
    var_EstadoCliente1 = EstadoCliente1.get()
    var_CEPCliente1 = CEPCliente1.get()
    var_NacionalidadeCliente1 = NacionalidadeCliente1.get()
    var_EstCivilCliente1 = cb_EstCivilCliente1.get()
    var_ProfissaoCliente1 = ProfissaoCliente1.get()
    var_EmailCliente1 = EmailCliente1.get()
    var_CelCliente1 = CelCliente1.get()

    Path(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '').mkdir(
        parents=True, exist_ok=True)

    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.underline = True
    font.color.rgb = RGBColor(0, 0, 255)
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
    paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.PARTES')
    paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.1 CONTRATADO:')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
    paragraph.add_run(
        'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.2 CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente1 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente1 + ', ' + var_EstCivilCliente1 + ', ' + var_ProfissaoCliente1 + ', '
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente1 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente1 + ','
                                                                                                                                                                                                                     ' residente e domiciliado(a) na ' + var_EndCliente1 + ', '
                                                                                                                                                                                                                                                                           'n° ' + var_NCliente1 + ', ' + var_BairroCliente1 + ' na cidade de ' + var_CidadeCliente1 + '/' + var_EstadoCliente1 + '. ').bold = False
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                       'residencial em “AUTOCAD”, conforme características do imóvel do CONTRATANTE e '
                                       'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                       + var_CidadeObra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                          'projeto até a liberação do Alvará. Para o imóvel: lote ' + var_LoteObra + ', quadra ' + var_QuadraObra + '; do loteamento '
                                                                                                                                                                    'denominado “' + var_BairroObra + '”, no município de ' + var_CidadeObra + '-SP.')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    font.italic = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('a)	Documentos necessários;\nb)	Livre acesso ao local.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3. VISITAS')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                       '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                       '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                       '3.4 Caso houver interesse do CONTRATANTE ')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.add_run('CONTRATANTE').bold = True
    paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                      'Valor para cada visita técnica é de R$ ' + var_Visita + ',00 (' + var_Visitaextenso + ') hora.')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
        '4.2 O valor deste contrato é de ')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    runner = paragraph.add_run("" + var_ValorContrato + ' (' + var_ValorContratoextenso + ')')
    runner.bold = True
    runner.underline = True

    paragraph.add_run(', que o ')
    runner = paragraph.add_run('CONTRATANTE ')
    runner.bold = True
    paragraph.add_run('se obriga a pagar ao  ')
    runner = paragraph.add_run('CONTRATADO ')
    runner.bold = True
    paragraph.add_run(
        'em ' + var_ParcelContrato + ',00 (' + var_ParcelContratoextenso + ') vezes mensais, com vencimento '
                                                                           'todo o dia ' + var_DiaVencimento + ' de cada mês, com início em '
        + var_InicioContrato + ', constituindo-se nenhuma '
                               'tolerância de qualquer recebimento depois do '
                               'prazo estipulado.\n\n 4.3 Ao ')
    runner = paragraph.add_run('CONTRATANTE ')
    paragraph.add_run(
        'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado ao ')
    runner = paragraph.add_run('CONTRATANTE ')
    runner.bold = True
    paragraph.add_run(
        ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5. MULTAS')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5.1')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run(' MULTA DE MORA: ')
    runner.bold = True
    paragraph.add_run('Fica estipulada a multa de ')
    runner = paragraph.add_run('10%')
    runner.bold = True
    paragraph.add_run(
        '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. O')
    runner = paragraph.add_run(' CONTRATANTE ')
    runner.bold = True
    paragraph.add_run(
        'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
        '6.2 Fica eleito o foro')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run(' HORTOLÂNDIA – SP')
    runner.bold = True
    paragraph.add_run(
        ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que o ')
    runner = paragraph.add_run('CONTRATANTE ')
    runner.bold = True
    paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                      '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')

    paragraph = document.add_paragraph(
        'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('' + var_CidadeObra + '/SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________          _____________________________________')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('CONTRATADO:                                            CONTRATANTE:')
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


    paragraph = document.add_paragraph('ROGÉRIO ROCHA SOARES                      ' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Contrato ' + var_NomeCliente1 + '.docx')

    # Memorial Descritivo
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                       'Local: ' + var_EndObra + ', nº ' + var_NObra + ' Lote: ' + var_LoteObra + ' – Quadra: ' + var_QuadraObra + '\n'
                                                                                                                                                   'Loteamento: ' + var_BairroObra + ' -  ' + var_CidadeObra + ' - SP\n'
                                                                                                                                                                                                               'Proprietário: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    runner = paragraph.add_run('' + var_NomeCliente1 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente1 + ' \n')
    runner.bold = True

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + var_AreaObra + ' m² ')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Descrição')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
    paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
    paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
    paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
    paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
    paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
    paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
    paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
    paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
    paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
    paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
    paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
    paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
    paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
    paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES                                                       Proprietário:' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'Engenheiro Civil                                                                            CPF:' + var_CPFCliente1 + '')
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'ART' + var_ArtObra + '')
    paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Memorial Descritivo ' + var_NomeCliente1 + '.docx')

    # RRC sem lei
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nEu')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(' ' + var_NomeCliente1 + ' ')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '_________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Do Proprietário')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente1 + '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra:' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Do Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Requerimento sem Lei_' + var_NomeCliente1 + '.docx')

    # RRC com lei
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\nEu')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(' ' + var_NomeCliente1 + ' ')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '_________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Do Proprietário')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente1 + '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra:' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Do Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Requerimento com Lei_' + var_NomeCliente1 + '.docx')

    # Procuração
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('PROCURAÇÃO ')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11.5)

    runner = paragraph.add_run('I - OUTORGANTE:')
    runner.bold = True

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente1 + ' CPF: ' + var_CPFCliente1 + '\n\n')

    runner = paragraph.add_run('II – OUTORGADO: ')
    runner.bold = True

    paragraph.add_run(
        '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

    runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
    runner.bold = True

    paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                      'Lote' + var_LoteObra + 'da Quadra ' + var_QuadraObra + ', localizado no endereço: '
                                                                              '' + var_EndObra + ' nº ' + var_NObra + ' Loteamento: ' + var_BairroObra + '.\n\n')

    runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
    runner.bold = True

    paragraph.add_run('\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                      '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

    runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + var_CidadeObra + ', ')
    runner.bold = True

    paragraph.add_run('especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('OUTORGANTE:')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('OUTORGANTE:')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'CPF: 183.125.858-77')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Procuração ' + var_NomeCliente1 + '.docx')

    # Declaração
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

    paragraph = document.add_paragraph('ANEXO I')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DECLARAÇÃO')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                       'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                       ' que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município'
                                       ' de Hortolândia”, eu,' + var_NomeCliente1 + ', (' + var_ProfissaoCliente1 + '), ')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run('Proprietário da obra ')
    runner.bold = True

    paragraph.add_run(
        'localizada à ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ',' \
                                                                                                                                                          'cidade de Hortolândia-SP, DECLARO estar ciente das disposições ' \
                                                                                                                                                          'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                          'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                          'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                          'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                          'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                          'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                          ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                          'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                          'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                          'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                          ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

    runner = paragraph.add_run('DECLARO ')
    runner.bold = True

    paragraph.add_run(
        'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

    runner = paragraph.add_run('um dos seguintes documentos: ')
    runner.bold = True

    paragraph = document.add_paragraph(
        '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run('aaaaaaaaaaa')
    font = runner.font
    font.color.rgb = RGBColor(255, 255, 255)

    paragraph.add_run('origem nativa;')

    paragraph = document.add_paragraph(
        '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('PROPRIETÁRIO')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)

    paragraph = document.add_paragraph('\nEm conformidade com o disposto no artigo 4º da '
                                       'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                       'que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município '
                                       'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                       'Autor do Projeto da obra localizada à '
                                       'Rua ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ',cidade de Hortolândia-SP,'
                                                                                                                                                                                ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                                                ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                                                'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(
        'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
    runner.underline = True

    paragraph.add_run(
        'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('AUTOR DO PROJETO')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)

    # runner_word.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Declaração ' + var_NomeCliente1 + '.docx')

    # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Obra: CONSTRUÇÃO RESIDENCIAL UNIFAMILIAR – R1')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('Local: ' + var_EndObra + '- N° ' + var_NObra + '')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Lote: ' + var_LoteObra + ' Quadra: ' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Loteamento: ' + var_BairroObra + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Município: ' + var_CidadeObra + '/SP')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Proprietário: ' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph(
        '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
    paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
    paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
    paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
    paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
    paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
    paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
    paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES                                        Proprietário: ' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'Engenheiro Civil                                                            CPF:' + var_CPFCliente1 + '')
    paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.style = document.styles.add_style('style31', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph(
        'SMPUGE: 1036/18')
    paragraph.style = document.styles.add_style('style32', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Memorial Descritivo Para Construção ' + var_NomeCliente1 + '.docx')

    # ---------------------Recibo---------------------------------------------------------------------------------------------
    document = Document()

    document.sections[0].header_distance = Cm(0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('RECIBO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                       '5070374192, recebi de ' + var_NomeCliente1 + ', '
                                                                                     'parte do pagamento para aprovação de projeto '
                                                                                     'arquitetônico a quantia de R$ ' + var_ValorParcelaContrato + ',00 (' + var_ValorParcelaContratoextenso +
                                       '), de um total de '
                                       'R$ ' + var_ValorContrato + ',00 (' + var_ValorContratoextenso + ').')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(12)

    paragraph = document.add_paragraph('\n\n' + var_CidadeObra + ' ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font.size = Pt(12)

    # footer section
    footer_section = document.sections[0]
    footer = footer_section.footer

    # footer text
    footer_text = footer.paragraphs[0]
    footer_text.text = "_______________________________________________________________________________________________" \
                       "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                       "\nCREA: 5070347192" \
                       "\nE-MAIL: rocha.soares@hotmail.com"
    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + '/Recibo ' + var_NomeCliente1 + '.docx')

    MsgBox = tk.messagebox.showwarning("SUCESSO", "DOCUMENTOS FEITO COM SUCESSO. NÃO ESQUEÇA DE PEGA-LOS",
                                       icon='question')
    if MsgBox == 'ok':
        root.deiconify()
        janela2.destroy()
        JanelaCliente1.destroy()


# Janela para fazer doc de dois clientes
def confirmar2():
    # variaveis da obra
    var_EndObra = endobra.get()
    var_NObra = NumObra.get()
    var_BairroObra = BairroObra.get()
    var_LoteObra = LoteObra.get()
    var_QuadraObra = QuadraObra.get()
    var_CidadeObra = CidadeObra.get()
    var_QuarteiraoObra = QuarteiraoObra.get()
    var_TipoObra = TipoObra.get()
    var_AreaObra = AreaObra.get()
    var_ArtObra = ArtObra.get()
    var_ValorContrato = ValorContrato.get()
    var_ValorContratoextenso = ValorContratoextenso.get()
    var_ParcelContrato = ParcelContrato.get()
    var_ParcelContratoextenso = ParcelContratoextenso.get()
    var_ValorParcelaContrato = ValorParcelaContrato.get()
    var_ValorParcelaContratoextenso = ValorParcelaContratoextenso.get()
    var_DiaVencimento = DiaVencimento.get()
    var_InicioContrato = InicioContrato.get()
    var_Visita = Visita.get()
    var_Visitaextenso = Visitaextenso.get()

    # variaveis do cliente 1
    var_NomeCliente1 = NomeCliente1.get()
    var_CPFCliente1 = CPFCliente1.get()
    var_RGCliente1 = RGCliente1.get()
    var_EndCliente1 = EndCliente1.get()
    var_NCliente1 = NCliente1.get()
    var_BairroCliente1 = BairroCliente1.get()
    var_CidadeCliente1 = CidadeCliente1.get()
    var_EstadoCliente1 = EstadoCliente1.get()
    var_CEPCliente1 = CEPCliente1.get()
    var_NacionalidadeCliente1 = NacionalidadeCliente1.get()
    var_EstCivilCliente1 = cb_EstCivilCliente1.get()
    var_ProfissaoCliente1 = ProfissaoCliente1.get()
    var_EmailCliente1 = EmailCliente1.get()
    var_CelCliente1 = CelCliente1.get()

    # variaveis do cliente 2
    var_NomeCliente2 = NomeCliente2.get()
    var_CPFCliente2 = CPFCliente2.get()
    var_RGCliente2 = RGCliente2.get()
    var_EndCliente2 = EndCliente2.get()
    var_NCliente2 = NCliente2.get()
    var_BairroCliente2 = BairroCliente2.get()
    var_CidadeCliente2 = CidadeCliente2.get()
    var_EstadoCliente2 = EstadoCliente2.get()
    var_CEPCliente2 = CEPCliente2.get()
    var_NacionalidadeCliente2 = NacionalidadeCliente2.get()
    var_EstCivilCliente2 = cb_EstCivilCliente2.get()
    var_ProfissaoCliente2 = ProfissaoCliente2.get()
    var_EmailCliente2 = EmailCliente2.get()
    var_CelCliente2 = CelCliente2.get()

    Path(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '').mkdir(
        parents=True, exist_ok=True)

    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.underline = True
    font.color.rgb = RGBColor(0, 0, 255)
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
    paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('1.PARTES')
    paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('1.1 CONTRATADO:')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
    paragraph.add_run(
        'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
    enter = document.add_paragraph('')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('1.1 CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente1 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente1 + ', ' + var_EstCivilCliente1 + ', ' + var_ProfissaoCliente1 + ', '
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente1 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente1 + ''
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente1 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente1 + ', ' + var_BairroCliente1 + ' na cidade de ' + var_CidadeCliente1 + '/' + var_EstadoCliente1 + '. ').bold = False
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('    CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente2 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente2 + ', ' + var_EstCivilCliente2 + ', ' + var_ProfissaoCliente2 + ', '
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente2 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente2 + ''
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente2 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente2 + ', ' + var_BairroCliente2 + ' na cidade de ' + var_CidadeCliente2 + '/' + var_EstadoCliente2 + '. ').bold = False
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    enter = document.add_paragraph('')


    paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                       'residencial em “AUTOCAD”, conforme características do imóvel dos CONTRATANTES e '
                                       'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                       + var_CidadeObra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                          'projeto até a liberação do Alvará. Para o imóvel: lote ' + var_LoteObra + ', quadra ' + var_QuadraObra + '; do loteamento '
                                                                                                                                                                    'denominado “' + var_BairroObra + '”, no município de ' + var_CidadeObra + '-SP.')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    font.italic = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('a)	Documentos necessários;\nb)	Livre acesso ao local.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')


    paragraph = document.add_paragraph('\n\n\n\n\n\n\n\n3. VISITAS')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                       '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                       '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                       '3.4 Caso houver interesse dos ')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.add_run('CONTRATANTES').bold = True
    paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                      'Valor para cada visita técnica é de R$ ' + var_Visita + ',00 (' + var_Visitaextenso + ') hora.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
        '4.2 O valor deste contrato é de ')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    runner = paragraph.add_run("" + var_ValorContrato + ' (' + var_ValorContratoextenso + ')')
    runner.bold = True
    runner.underline = True

    paragraph.add_run(', que o ')
    runner = paragraph.add_run('CONTRATANTE ')
    runner.bold = True
    paragraph.add_run('se obriga a pagar ao  ')
    runner = paragraph.add_run('CONTRATADO ')
    runner.bold = True
    paragraph.add_run(
        'em ' + var_ParcelContrato + ',00 (' + var_ParcelContratoextenso + ') vezes mensais, com vencimento '
                                                                           'todo o dia ' + var_DiaVencimento + ' de cada mês, com início em '
        + var_InicioContrato + ', constituindo-se nenhuma '
                               'tolerância de qualquer recebimento depois do '
                               'prazo estipulado.\n\n 4.3 Ao ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado aos ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('\n\n\n\n\n\n\n\n\n5. MULTAS')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5.1 ')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run('MULTA DE MORA: ')
    runner.bold = True
    paragraph.add_run('Fica estipulada a multa de ')
    runner = paragraph.add_run('10%')
    runner.bold = True
    paragraph.add_run(
        '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. Os')
    runner = paragraph.add_run(' CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
        '6.2 Fica eleito o foro')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run(' HORTOLÂNDIA – SP')
    runner.bold = True
    paragraph.add_run(
        ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que os ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                      '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')

    paragraph = document.add_paragraph(
        'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('' + var_CidadeObra + ' / SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________          _____________________________________')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('CONTRATADO:                                            CONTRATANTE:')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('ROGÉRIO ROCHA SOARES                      ' + var_NomeCliente1 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________')
    paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('CONTRATANTE:')
    paragraph.style = document.styles.add_style('style18.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph('' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Contrato ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # Memorial Descritivo
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                       'Local: ' + var_EndObra + ', nº ' + var_NObra + ' Lote: ' + var_LoteObra + ' – Quadra: ' + var_QuadraObra + '\n'
                                                                                                                                                   'Loteamento: ' + var_BairroObra + ' -  ' + var_CidadeObra + ' - SP\n'
                                                                                                                                                                                                               'Proprietário: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    runner = paragraph.add_run('' + var_NomeCliente1 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente1 + ' \n')
    runner.bold = True

    runner = paragraph.add_run('                     ' + var_NomeCliente2 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente2 + ' \n')
    runner.bold = True

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + var_AreaObra + ' m² ')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Descrição')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede pública de esgoto.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário: ' + var_NomeCliente1 + '                                                Proprietário: ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '                                                                        CPF:' + var_CPFCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Engenheiro Civil')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style20.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style21.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ART' + var_ArtObra + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style22.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Memorial Descritivo ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # RRC sem lei
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nNós,')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(' ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + ' ')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    paragraph.paragraph_format.space_after = Pt(0)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dos Proprietários')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: ' + var_NomeCliente1 + ' e '+var_NomeCliente2+'\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado: ' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '\nEmail: '+var_EmailCliente1+'')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra: ' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('\nDo Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Requerimento sem Lei ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # RRC com lei
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nNós,')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(' ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + ' ')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dos Proprietários')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: ' + var_NomeCliente1 + ' e '+var_NomeCliente2+  '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N° ' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento: ' + var_LoteObra + '\n'
                                                                                                                                                            'CEP: ' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado: ' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '\nEmail: '+var_EmailCliente1+'')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra: ' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('\nDo Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Requerimento com Lei ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # Procuração
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('PROCURAÇÃO ')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11.5)

    runner = paragraph.add_run('I - OUTORGANTES:')
    runner.bold = True

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente1 + ' CPF: ' + var_CPFCliente1 + '')

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente2 + ' CPF: ' + var_CPFCliente2 + '\n\n')

    runner = paragraph.add_run('II – OUTORGADO: ')
    runner.bold = True

    paragraph.add_run(
        '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

    runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
    runner.bold = True

    paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                      'Lote ' + var_LoteObra + ' da Quadra ' + var_QuadraObra + ', localizado no endereço: '
                                                                              '' + var_EndObra + ' nº ' + var_NObra + ' Loteamento: ' + var_BairroObra + '.\n\n')

    runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
    runner.bold = True

    paragraph.add_run('\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                      '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

    runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + var_CidadeObra + ', ')
    runner.bold = True

    paragraph.add_run('especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n\n_________________________________                            _________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        'OUTORGANTE:                                                                                                 OUTORGANTE:')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n\n\n\n_________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('OUTORGADO:')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)

    paragraph = document.add_paragraph(
        'CPF: 183.125.858-77')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Procuração ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # Declaração
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

    paragraph = document.add_paragraph('ANEXO I')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DECLARAÇÃO')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                       'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                       ' que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município'
                                       ' de Hortolândia”, nós, ' + var_NomeCliente1 + ', (' + var_ProfissaoCliente1 + ') e ' + var_NomeCliente2 + ', (' + var_ProfissaoCliente2 + '),')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(' Proprietários da obra ')
    runner.bold = True

    paragraph.add_run(
        'localizada à ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ', ' \
                                                                                                                                                          'cidade de Hortolândia-SP, DECLARAMOS estar ciente das disposições ' \
                                                                                                                                                          'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                          'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                          'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                          'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                          'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                          'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                          ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                          'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                          'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                          'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                          ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

    runner = paragraph.add_run('DECLARAMOS ')
    runner.bold = True

    paragraph.add_run(
        'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

    runner = paragraph.add_run('um dos seguintes documentos: ')
    runner.bold = True

    paragraph = document.add_paragraph(
        '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run('aaaaaaaaaaa')
    font = runner.font
    font.color.rgb = RGBColor(255, 255, 255)

    paragraph.add_run('origem nativa;')

    paragraph = document.add_paragraph(
        '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_____________________________________')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente1 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_____________________________________')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style8.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style9.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('\nEm conformidade com o disposto no artigo 4º da '
                                       'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                       'que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município '
                                       'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                       'Autor do Projeto da obra localizada à '
                                       'Rua ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ', cidade de Hortolândia-SP,'
                                                                                                                                                                                ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                                                ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                                                'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(
        'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
    runner.underline = True

    paragraph.add_run(
        'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('AUTOR DO PROJETO')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # runner_word.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Declaração ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Obra: CONSTRUÇÃO RESIDENCIAL UNIFAMILIAR – R1')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    paragraph = document.add_paragraph('Local: ' + var_EndObra + '- N° ' + var_NObra + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Lote: ' + var_LoteObra + ' Quadra: ' + var_QuadraObra + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Loteamento: ' + var_BairroObra + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Município: ' + var_CidadeObra + '/SP')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Proprietário(s): ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph(
        '\n01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                    _____________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário: ' + var_NomeCliente1 + '                               Proprietário: ' + var_NomeCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '                                                        CPF:' + var_CPFCliente2 + '')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________')
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style28.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style29.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Engenheiro Civil')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style30.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style31.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'SMPUGE: 1036/18')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = document.styles.add_style('style32.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Memorial Descritivo Para Construção ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    # ---------------------Recibo---------------------------------------------------------------------------------------------
    document = Document()

    document.sections[0].header_distance = Cm(0.5)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('RECIBO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                       '5070374192, recebi de ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + ', '
                                                                                                                'parte do pagamento para aprovação de projeto '
                                                                                                                'arquitetônico a quantia de R$ ' + var_ValorParcelaContrato + ',00 (' + var_ValorParcelaContratoextenso +
                                       '), de um total de '
                                       'R$ ' + var_ValorContrato + ',00 (' + var_ValorContratoextenso + ').')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(12)

    paragraph = document.add_paragraph('\n\n' + var_CidadeObra + ', ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font.size = Pt(12)

    # footer section
    footer_section = document.sections[0]
    footer = footer_section.footer

    # footer text
    footer_text = footer.paragraphs[0]
    footer_text.text = "_______________________________________________________________________________________________" \
                       "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                       "\nCREA: 5070347192" \
                       "\nE-MAIL: rocha.soares@hotmail.com"
    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '/Recibo ' + var_NomeCliente1 + ' e ' + var_NomeCliente2 + '.docx')

    MsgBox = tk.messagebox.showwarning("SUCESSO", "DOCUMENTOS FEITO COM SUCESSO. NÃO ESQUEÇA DE PEGA-LOS",
                                       icon='question')
    if MsgBox == 'ok':
        root.deiconify()
        janela2.destroy()
        JanelaCliente2.destroy()


# Janela para fazer doc de três clientes
def confirmar3():
    # variaveis da obra
    var_EndObra = endobra.get()
    var_NObra = NumObra.get()
    var_BairroObra = BairroObra.get()
    var_LoteObra = LoteObra.get()
    var_QuadraObra = QuadraObra.get()
    var_CidadeObra = CidadeObra.get()
    var_QuarteiraoObra = QuarteiraoObra.get()
    var_TipoObra = TipoObra.get()
    var_AreaObra = AreaObra.get()
    var_ArtObra = ArtObra.get()
    var_ValorContrato = ValorContrato.get()
    var_ValorContratoextenso = ValorContratoextenso.get()
    var_ParcelContrato = ParcelContrato.get()
    var_ParcelContratoextenso = ParcelContratoextenso.get()
    var_ValorParcelaContrato = ValorParcelaContrato.get()
    var_ValorParcelaContratoextenso = ValorParcelaContratoextenso.get()
    var_DiaVencimento = DiaVencimento.get()
    var_InicioContrato = InicioContrato.get()
    var_Visita = Visita.get()
    var_Visitaextenso = Visitaextenso.get()

    # variaveis do cliente 1
    var_NomeCliente1 = NomeCliente1.get()
    var_CPFCliente1 = CPFCliente1.get()
    var_RGCliente1 = RGCliente1.get()
    var_EndCliente1 = EndCliente1.get()
    var_NCliente1 = NCliente1.get()
    var_BairroCliente1 = BairroCliente1.get()
    var_CidadeCliente1 = CidadeCliente1.get()
    var_EstadoCliente1 = EstadoCliente1.get()
    var_CEPCliente1 = CEPCliente1.get()
    var_NacionalidadeCliente1 = NacionalidadeCliente1.get()
    var_EstCivilCliente1 = cb_EstCivilCliente1.get()
    var_ProfissaoCliente1 = ProfissaoCliente1.get()
    var_EmailCliente1 = EmailCliente1.get()
    var_CelCliente1 = CelCliente1.get()

    # variaveis do cliente 2
    var_NomeCliente2 = NomeCliente2.get()
    var_CPFCliente2 = CPFCliente2.get()
    var_RGCliente2 = RGCliente2.get()
    var_EndCliente2 = EndCliente2.get()
    var_NCliente2 = NCliente2.get()
    var_BairroCliente2 = BairroCliente2.get()
    var_CidadeCliente2 = CidadeCliente2.get()
    var_EstadoCliente2 = EstadoCliente2.get()
    var_CEPCliente2 = CEPCliente2.get()
    var_NacionalidadeCliente2 = NacionalidadeCliente2.get()
    var_EstCivilCliente2 = cb_EstCivilCliente2.get()
    var_ProfissaoCliente2 = ProfissaoCliente2.get()
    var_EmailCliente2 = EmailCliente2.get()
    var_CelCliente2 = CelCliente2.get()

    # variaveis do cliente 3
    var_NomeCliente3 = NomeCliente3.get()
    var_CPFCliente3 = CPFCliente3.get()
    var_RGCliente3 = RGCliente3.get()
    var_EndCliente3 = EndCliente3.get()
    var_NCliente3 = NCliente3.get()
    var_BairroCliente3 = BairroCliente3.get()
    var_CidadeCliente3 = CidadeCliente3.get()
    var_EstadoCliente3 = EstadoCliente3.get()
    var_CEPCliente3 = CEPCliente3.get()
    var_NacionalidadeCliente3 = NacionalidadeCliente3.get()
    var_EstCivilCliente3 = cb_EstCivilCliente3.get()
    var_ProfissaoCliente3 = ProfissaoCliente3.get()
    var_EmailCliente3 = EmailCliente3.get()
    var_CelCliente3 = CelCliente3.get()

    Path(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '').mkdir(
        parents=True, exist_ok=True)

    document = Document()

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.underline = True
    font.color.rgb = RGBColor(0, 0, 255)
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
    paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.PARTES')
    paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.1 CONTRATADO:')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
    paragraph.add_run(
        'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.1 CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente1 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente1 + ', ' + var_EstCivilCliente1 + ', ' + var_ProfissaoCliente1 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente1 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente1 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente1 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente1 + ', ' + var_BairroCliente1 + ' na cidade de ' + var_CidadeCliente1 + '/' + var_EstadoCliente1 + '. ').bold = False

    paragraph = document.add_paragraph('    CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente2 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente2 + ', ' + var_EstCivilCliente2 + ', ' + var_ProfissaoCliente2 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente2 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente2 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente2 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente2 + ', ' + var_BairroCliente2 + ' na cidade de ' + var_CidadeCliente2 + '/' + var_EstadoCliente2 + '. ').bold = False

    paragraph = document.add_paragraph('    CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente3 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente3 + ', ' + var_EstCivilCliente3 + ', ' + var_ProfissaoCliente3 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente3 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente3 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente3 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente3 + ', ' + var_BairroCliente3 + ' na cidade de ' + var_CidadeCliente3 + '/' + var_EstadoCliente3 + '. ').bold = False

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                       'residencial em “AUTOCAD”, conforme características do imóvel dos CONTRATANTES e '
                                       'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                       + var_CidadeObra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                          'projeto até a liberação do Alvará. Para o imóvel: lote ' + var_LoteObra + ', quadra ' + var_QuadraObra + '; do loteamento '
                                                                                                                                                                    'denominado “' + var_BairroObra + '”, no município de ' + var_CidadeObra + '-SP.')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    font.italic = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('a)	Documentos necessários;\nb)	Livre acesso ao local.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3. VISITAS')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                       '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                       '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                       '3.4 Caso houver interesse dos ')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.add_run('CONTRATANTES').bold = True
    paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                      'Valor para cada visita técnica é de R$ ' + var_Visita + ',00 (' + var_Visitaextenso + ') hora.')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
        '4.2 O valor deste contrato é de')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    runner = paragraph.add_run("" + var_ValorContrato + ' (' + var_ValorContratoextenso + ')')
    runner.bold = True
    runner.underline = True

    paragraph.add_run(', que o ')
    runner = paragraph.add_run('CONTRATANTE ')
    runner.bold = True
    paragraph.add_run('se obriga a pagar ao  ')
    runner = paragraph.add_run('CONTRATADO ')
    runner.bold = True
    paragraph.add_run(
        'em ' + var_ParcelContrato + ',00 (' + var_ParcelContratoextenso + ') vezes mensais, com vencimento '
                                                                           'todo o dia ' + var_DiaVencimento + ' de cada mês, com início em '
        + var_InicioContrato + ', constituindo-se nenhuma '
                               'tolerância de qualquer recebimento depois do '
                               'prazo estipulado.\n\n 4.3 Ao ')
    runner = paragraph.add_run('CONTRATANTES ')
    paragraph.add_run(
        'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado aos ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5. MULTAS')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5.1')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run('MULTA DE MORA: ')
    runner.bold = True
    paragraph.add_run('Fica estipulada a multa de ')
    runner = paragraph.add_run('10%')
    runner.bold = True
    paragraph.add_run(
        '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. Os')
    runner = paragraph.add_run(' CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
        '6.2 Fica eleito o foro')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run(' HORTOLÂNDIA – SP')
    runner.bold = True
    paragraph.add_run(
        ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que os ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                      '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('' + var_CidadeObra + '/SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________          _____________________________________')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True

    paragraph = document.add_paragraph('CONTRATADO:                                            CONTRATANTE:')
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    paragraph = document.add_paragraph('ROGÉRIO ROCHA SOARES                      ' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________          _____________________________________')
    paragraph.style = document.styles.add_style('style17.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True

    paragraph = document.add_paragraph('CONTRATANTE:                                            CONTRATANTE:')
    paragraph.style = document.styles.add_style('style18.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    paragraph = document.add_paragraph('' + var_NomeCliente2 + '                      ' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style19.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Contrato ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # Memorial Descritivo
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                       'Local: ' + var_EndObra + ', nº ' + var_NObra + ' Lote: ' + var_LoteObra + ' – Quadra: ' + var_QuadraObra + '\n'
                                                                                                                                                   'Loteamento: ' + var_BairroObra + ' -  ' + var_CidadeObra + ' - SP\n'
                                                                                                                                                                                                               'Proprietário: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    runner = paragraph.add_run('' + var_NomeCliente1 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente1 + ' \n')
    runner.bold = True

    runner = paragraph.add_run('                     ' + var_NomeCliente2 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente2 + ' \n')
    runner.bold = True

    runner = paragraph.add_run('                     ' + var_NomeCliente3 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente3 + ' \n')
    runner.bold = True

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + var_AreaObra + ' m² ')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Descrição')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
    paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
    paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
    paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
    paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
    paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
    paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
    paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
    paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
    paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
    paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
    paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
    paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
    paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
    paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário:' + var_NomeCliente1 + '                                                Proprietário:' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '                                                                        CPF:' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário:' + var_NomeCliente3 + '                                                 ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente3 + '                                                                         Engenheiro Civil')
    paragraph.style = document.styles.add_style('style20.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '                                                                     CREA: 5070347192-SP')
    paragraph.style = document.styles.add_style('style21.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '                                                                     ART' + var_ArtObra + '')
    paragraph.style = document.styles.add_style('style22.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Memorial Descritivo ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # RRC sem lei
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nNós,')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(' ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente2 + '')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________ ')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style6.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente3 + '')
    paragraph.style = document.styles.add_style('style7.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dos Proprietários')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente1 + '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente2 + '\n'
                                                                    'Endereço: ' + var_EndCliente2 + ' N°' + var_NCliente2 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente2 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente2 + '-' + var_EstadoCliente2 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente2 + '')
    paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente3 + '\n'
                                                                    'Endereço: ' + var_EndCliente3 + ' N°' + var_NCliente3 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente3 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente3 + '-' + var_EstadoCliente3 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente3 + '')
    paragraph.style = document.styles.add_style('style10.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra:' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Requerimento sem Lei ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # RRC com lei
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nNós,')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(' ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente2 + '')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________ ')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style6.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente3 + '')
    paragraph.style = document.styles.add_style('style7.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dos Proprietários')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente1 + '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente2 + '\n'
                                                                    'Endereço: ' + var_EndCliente2 + ' N°' + var_NCliente2 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente2 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente2 + '-' + var_EstadoCliente2 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente2 + '')
    paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra:' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Requerimento com Lei ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # Procuração
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('PROCURAÇÃO ')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11.5)

    runner = paragraph.add_run('I - OUTORGANTES:')
    runner.bold = True

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente1 + ' CPF: ' + var_CPFCliente1 + '')

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente2 + ' CPF: ' + var_CPFCliente2 + '')

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente3 + ' CPF: ' + var_CPFCliente3 + '\n\n')

    runner = paragraph.add_run('II – OUTORGADO: ')
    runner.bold = True

    paragraph.add_run(
        '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

    runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
    runner.bold = True

    paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                      'Lote' + var_LoteObra + 'da Quadra ' + var_QuadraObra + ', localizado no endereço: '
                                                                              '' + var_EndObra + ' nº ' + var_NObra + ' Loteamento: ' + var_BairroObra + '.\n\n')

    runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
    runner.bold = True

    paragraph.add_run('\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                      '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

    runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + var_CidadeObra + ',')
    runner.bold = True

    paragraph.add_run('especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        'OUTORGANTE:                                                                                                OUTORGANTE:')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style4.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        'OUTORGANTE:')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style6.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente3 + '')
    paragraph.style = document.styles.add_style('style7.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n\_________________________________')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('OUTORGADO:')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)

    paragraph = document.add_paragraph(
        'CPF: 183.125.858-77')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Procuração ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # Declaração
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(3.49)
        section.bottom_margin = Cm(1.1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

    paragraph = document.add_paragraph('ANEXO I')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DECLARAÇÃO')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                       'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                       ' que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município'
                                       ' de Hortolândia”, nós, ' + var_NomeCliente1 + ', (' + var_ProfissaoCliente1 + '), ' + var_NomeCliente2 + ', (' + var_ProfissaoCliente2 + ') e ' + var_NomeCliente3 + ', (' + var_ProfissaoCliente3 + '),')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(' Proprietários da obra ')
    runner.bold = True

    paragraph.add_run(
        'localizada à ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ',' \
                                                                                                                                                          'cidade de Hortolândia-SP, DECLARAMOS estar ciente das disposições ' \
                                                                                                                                                          'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                          'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                          'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                          'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                          'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                          'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                          ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                          'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                          'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                          'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                          ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

    runner = paragraph.add_run('DECLARAMOS ')
    runner.bold = True

    paragraph.add_run(
        'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

    runner = paragraph.add_run('um dos seguintes documentos: ')
    runner.bold = True

    paragraph = document.add_paragraph(
        '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run('aaaaaaaaaaa')
    font = runner.font
    font.color.rgb = RGBColor(255, 255, 255)

    paragraph.add_run('origem nativa;')

    paragraph = document.add_paragraph(
        '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style9.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style9.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('\n\n\n\n\n\n\n\nEm conformidade com o disposto no artigo 4º da '
                                       'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                       'que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município '
                                       'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                       'Autor do Projeto da obra localizada à '
                                       'Rua ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ',cidade de Hortolândia-SP,'
                                                                                                                                                                                ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                                                ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                                                'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(
        'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
    runner.underline = True

    paragraph.add_run(
        'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('AUTOR DO PROJETO')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # runner_word.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Declaração ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Obra: REGULARIZAÇÃO E CONSTRUÇÃO RESIDENCIAL MULTIFAMILIAR – R2')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('(DE ACORDO COM A LEI 3.491/2018 - ANISTIA)')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('Local: ' + var_EndObra + '- N° ' + var_NObra + '')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Lote: ' + var_LoteObra + ' Quadra: ' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Loteamento: ' + var_LoteObra + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Município: ' + var_CidadeObra + '/SP')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph(
        'Proprietário(s): ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph(
        '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
    paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
    paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
    paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
    paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
    paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
    paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
    paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário: ' + var_NomeCliente1 + '                               Proprietário: ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '                                                        CPF:' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style28.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário: ' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style29.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente3 + '')
    paragraph.style = document.styles.add_style('style30.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style28.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style29.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Engenheiro Civil')
    paragraph.style = document.styles.add_style('style30.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.style = document.styles.add_style('style31.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'SMPUGE: 1036/18')
    paragraph.style = document.styles.add_style('style32.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Memorial Descritivo Para Construção ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    # ---------------------Recibo---------------------------------------------------------------------------------------------
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('RECIBO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                       '5070374192, recebi de ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + ', '
                                                                                                                                          'parte do pagamento para aprovação de projeto '
                                                                                                                                          'arquitetônico a quantia de R$ ' + var_ValorParcelaContrato + ''
                                                                                                                                                                                                        ', de um total de '
                                                                                                                                                                                                        'R$ ' + var_ValorContrato + '.')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(12)

    paragraph = document.add_paragraph('\n\n' + var_CidadeObra + ' ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font.size = Pt(12)

    # footer section
    footer_section = document.sections[0]
    footer = footer_section.footer

    # footer text
    footer_text = footer.paragraphs[0]
    footer_text.text = "_______________________________________________________________________________________________" \
                       "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                       "\nCREA: 5070347192" \
                       "\nE-MAIL: rocha.soares@hotmail.com"
    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '/Recibo ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ' e ' + var_NomeCliente3 + '.docx')

    MsgBox = tk.messagebox.showwarning("SUCESSO", "DOCUMENTOS FEITO COM SUCESSO. NÃO ESQUEÇA DE PEGA-LOS",
                                       icon='question')
    if MsgBox == 'ok':
        root.deiconify()
        janela2.destroy()
        JanelaCliente3.destroy()


# Janela para fazer doc de quatro clientes
def confirmar4():
    # variaveis da obra
    var_EndObra = endobra.get()
    var_NObra = NumObra.get()
    var_BairroObra = BairroObra.get()
    var_LoteObra = LoteObra.get()
    var_QuadraObra = QuadraObra.get()
    var_CidadeObra = CidadeObra.get()
    var_QuarteiraoObra = QuarteiraoObra.get()
    var_TipoObra = TipoObra.get()
    var_AreaObra = AreaObra.get()
    var_ArtObra = ArtObra.get()
    var_ValorContrato = ValorContrato.get()
    var_ValorContratoextenso = ValorContratoextenso.get()
    var_ParcelContrato = ParcelContrato.get()
    var_ParcelContratoextenso = ParcelContratoextenso.get()
    var_ValorParcelaContrato = ValorParcelaContrato.get()
    var_ValorParcelaContratoextenso = ValorParcelaContratoextenso.get()
    var_DiaVencimento = DiaVencimento.get()
    var_InicioContrato = InicioContrato.get()
    var_Visita = Visita.get()
    var_Visitaextenso = Visitaextenso.get()

    # variaveis do cliente 1
    var_NomeCliente1 = NomeCliente1.get()
    var_CPFCliente1 = CPFCliente1.get()
    var_RGCliente1 = RGCliente1.get()
    var_EndCliente1 = EndCliente1.get()
    var_NCliente1 = NCliente1.get()
    var_BairroCliente1 = BairroCliente1.get()
    var_CidadeCliente1 = CidadeCliente1.get()
    var_EstadoCliente1 = EstadoCliente1.get()
    var_CEPCliente1 = CEPCliente1.get()
    var_NacionalidadeCliente1 = NacionalidadeCliente1.get()
    var_EstCivilCliente1 = cb_EstCivilCliente1.get()
    var_ProfissaoCliente1 = ProfissaoCliente1.get()
    var_EmailCliente1 = EmailCliente1.get()
    var_CelCliente1 = CelCliente1.get()

    # variaveis do cliente 2
    var_NomeCliente2 = NomeCliente2.get()
    var_CPFCliente2 = CPFCliente2.get()
    var_RGCliente2 = RGCliente2.get()
    var_EndCliente2 = EndCliente2.get()
    var_NCliente2 = NCliente2.get()
    var_BairroCliente2 = BairroCliente2.get()
    var_CidadeCliente2 = CidadeCliente2.get()
    var_EstadoCliente2 = EstadoCliente2.get()
    var_CEPCliente2 = CEPCliente2.get()
    var_NacionalidadeCliente2 = NacionalidadeCliente2.get()
    var_EstCivilCliente2 = cb_EstCivilCliente2.get()
    var_ProfissaoCliente2 = ProfissaoCliente2.get()
    var_EmailCliente2 = EmailCliente2.get()
    var_CelCliente2 = CelCliente2.get()

    # variaveis do cliente 3
    var_NomeCliente3 = NomeCliente3.get()
    var_CPFCliente3 = CPFCliente3.get()
    var_RGCliente3 = RGCliente3.get()
    var_EndCliente3 = EndCliente3.get()
    var_NCliente3 = NCliente3.get()
    var_BairroCliente3 = BairroCliente3.get()
    var_CidadeCliente3 = CidadeCliente3.get()
    var_EstadoCliente3 = EstadoCliente3.get()
    var_CEPCliente3 = CEPCliente3.get()
    var_NacionalidadeCliente3 = NacionalidadeCliente3.get()
    var_EstCivilCliente3 = cb_EstCivilCliente3.get()
    var_ProfissaoCliente3 = ProfissaoCliente3.get()
    var_EmailCliente3 = EmailCliente3.get()
    var_CelCliente3 = CelCliente3.get()

    # variaveis do cliente 4
    var_NomeCliente4 = NomeCliente4.get()
    var_CPFCliente4 = CPFCliente4.get()
    var_RGCliente4 = RGCliente4.get()
    var_EndCliente4 = EndCliente4.get()
    var_NCliente4 = NCliente4.get()
    var_BairroCliente4 = BairroCliente4.get()
    var_CidadeCliente4 = CidadeCliente4.get()
    var_EstadoCliente4 = EstadoCliente4.get()
    var_CEPCliente4 = CEPCliente4.get()
    var_NacionalidadeCliente4 = NacionalidadeCliente4.get()
    var_EstCivilCliente4 = cb_EstCivilCliente4.get()
    var_ProfissaoCliente4 = ProfissaoCliente4.get()
    var_EmailCliente4 = EmailCliente4.get()
    var_CelCliente4 = CelCliente4.get()

    Path(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '').mkdir(
        parents=True, exist_ok=True)

    document = Document()

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.underline = True
    font.color.rgb = RGBColor(0, 0, 255)
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
    paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.PARTES')
    paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.1 CONTRATADO:')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
    paragraph.add_run(
        'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('1.1 CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente1 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente1 + ', ' + var_EstCivilCliente1 + ', ' + var_ProfissaoCliente1 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente1 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente1 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente1 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente1 + ', ' + var_BairroCliente1 + ' na cidade de ' + var_CidadeCliente1 + '/' + var_EstadoCliente1 + '. ').bold = False

    paragraph = document.add_paragraph('    CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente2 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente2 + ', ' + var_EstCivilCliente2 + ', ' + var_ProfissaoCliente2 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente2 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente2 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente2 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente2 + ', ' + var_BairroCliente2 + ' na cidade de ' + var_CidadeCliente2 + '/' + var_EstadoCliente2 + '. ').bold = False

    paragraph = document.add_paragraph('    CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente3 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente3 + ', ' + var_EstCivilCliente3 + ', ' + var_ProfissaoCliente3 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente3 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente3 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente3 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente3 + ', ' + var_BairroCliente3 + ' na cidade de ' + var_CidadeCliente3 + '/' + var_EstadoCliente3 + '. ').bold = False

    paragraph = document.add_paragraph('    CONTRATANTE: ')
    paragraph.style = document.styles.add_style('style2.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    paragraph.add_run('' + var_NomeCliente4 + ', ').underline = True
    paragraph.add_run('' + var_NacionalidadeCliente4 + ', ' + var_EstCivilCliente4 + ', ' + var_ProfissaoCliente4 + ','
                                                                                                                    'portador(a) do RG n° ' + var_RGCliente4 + ' SSP/SP, inscrito(a) no CPF n° ' + var_CPFCliente4 + ' '
                                                                                                                                                                                                                     ', residente e domiciliado(a) na ' + var_EndCliente4 + ', '
                                                                                                                                                                                                                                                                            'n° ' + var_NCliente4 + ', ' + var_BairroCliente4 + ' na cidade de ' + var_CidadeCliente4 + '/' + var_EstadoCliente4 + '. ').bold = False

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                       'residencial em “AUTOCAD”, conforme características do imóvel dos CONTRATANTES e '
                                       'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                       + var_CidadeObra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                          'projeto até a liberação do Alvará. Para o imóvel: lote ' + var_LoteObra + ', quadra ' + var_QuadraObra + '; do loteamento '
                                                                                                                                                                    'denominado “' + var_BairroObra + '”, no município de ' + var_CidadeObra + '-SP.')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    font.italic = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('a)	Documentos necessários;\nb)	Livre acesso ao local.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3. VISITAS')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                       '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                       '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                       '3.4 Caso houver interesse dos ')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.add_run('CONTRATANTES').bold = True
    paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                      'Valor para cada visita técnica é de R$ ' + var_Visita + ',00 (' + var_Visitaextenso + ') hora.')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
        '4.2 O valor deste contrato é de')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    runner = paragraph.add_run("" + var_ValorContrato + ' (' + var_ValorContratoextenso + ')')
    runner.bold = True
    runner.underline = True

    paragraph.add_run(', que o ')
    runner = paragraph.add_run('CONTRATANTE ')
    runner.bold = True
    paragraph.add_run('se obriga a pagar ao  ')
    runner = paragraph.add_run('CONTRATADO ')
    runner.bold = True
    paragraph.add_run(
        'em ' + var_ParcelContrato + ',00 (' + var_ParcelContratoextenso + ') vezes mensais, com vencimento '
                                                                           'todo o dia ' + var_DiaVencimento + ' de cada mês, com início em '
        + var_InicioContrato + ', constituindo-se nenhuma '
                               'tolerância de qualquer recebimento depois do '
                               'prazo estipulado.\n\n 4.3 Ao ')
    runner = paragraph.add_run('CONTRATANTES ')
    paragraph.add_run(
        'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado aos ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5. MULTAS')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('5.1')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run('MULTA DE MORA: ')
    runner.bold = True
    paragraph.add_run('Fica estipulada a multa de ')
    runner = paragraph.add_run('10%')
    runner.bold = True
    paragraph.add_run(
        '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. Os')
    runner = paragraph.add_run(' CONTRATANTES ')
    runner.bold = True
    paragraph.add_run(
        'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
        '6.2 Fica eleito o foro')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    runner = paragraph.add_run(' HORTOLÂNDIA – SP')
    runner.bold = True
    paragraph.add_run(
        ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que os ')
    runner = paragraph.add_run('CONTRATANTES ')
    runner.bold = True
    paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                      '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'

    paragraph = document.add_paragraph('' + var_CidadeObra + '/SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________          _____________________________________')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True

    paragraph = document.add_paragraph('CONTRATANTE:                                            CONTRATANTE:')
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    paragraph = document.add_paragraph('' + var_NomeCliente1 + '                      ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    enter = document.add_paragraph('')
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        '_________________________________          _____________________________________')
    paragraph.style = document.styles.add_style('style17.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True

    paragraph = document.add_paragraph('CONTRATANTE:                                            CONTRATANTE:')
    paragraph.style = document.styles.add_style('style18.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    paragraph = document.add_paragraph('' + var_NomeCliente3 + '                      ' + var_NomeCliente4 + '')
    paragraph.style = document.styles.add_style('style19.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    paragraph = document.add_paragraph(
        '_________________________________')
    paragraph.style = document.styles.add_style('style17.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.italic = True
    font.bold = True

    paragraph = document.add_paragraph('CONTRATADO:')
    paragraph.style = document.styles.add_style('style18.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    paragraph = document.add_paragraph('ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style194.', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    font.bold = True

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Contrato ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '.docx')

    # Memorial Descritivo
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                       'Local: ' + var_EndObra + ', nº ' + var_NObra + ' Lote: ' + var_LoteObra + ' – Quadra: ' + var_QuadraObra + '\n'
                                                                                                                                                   'Loteamento: ' + var_BairroObra + ' -  ' + var_CidadeObra + ' - SP\n'
                                                                                                                                                                                                               'Proprietário: ')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    runner = paragraph.add_run('' + var_NomeCliente1 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente1 + ' \n')
    runner.bold = True

    runner = paragraph.add_run('                     ' + var_NomeCliente2 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente2 + ' \n')
    runner.bold = True

    runner = paragraph.add_run('                     ' + var_NomeCliente3 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente3 + ' \n')
    runner.bold = True

    runner = paragraph.add_run('                     ' + var_NomeCliente4 + ' \n')
    runner.bold = True
    paragraph.add_run('                     CPF:')
    runner = paragraph.add_run('' + var_CPFCliente4 + ' \n')
    runner.bold = True

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + var_AreaObra + ' m² ')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)

    enter = document.add_paragraph('')

    paragraph = document.add_paragraph('Descrição')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    enter = document.add_paragraph('')

    paragraph = document.add_paragraph(
        'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
    paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
    paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
    paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
    paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
    paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
    paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
    paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
    paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
    paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
    paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
    paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
    paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
    paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
    paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário:' + var_NomeCliente1 + '                                                Proprietário:' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '                                                                        CPF:' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário:' + var_NomeCliente3 + '                                                 Proprietário:' + var_NomeCliente4 + '')
    paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente3 + '                                                                         CPF:' + var_CPFCliente4 + '')
    paragraph.style = document.styles.add_style('style20.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                                    _____________________________________')
    paragraph.style = document.styles.add_style('style17.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style19.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Engenheiro Civil')
    paragraph.style = document.styles.add_style('style20.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.style = document.styles.add_style('style21.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ART' + var_ArtObra + '')
    paragraph.style = document.styles.add_style('style22.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Memorial Descritivo ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '.docx')

    # RRC sem lei
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nNós,')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(
        ' ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dos Proprietários')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente1 + '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente2 + '\n'
                                                                    'Endereço: ' + var_EndCliente2 + ' N°' + var_NCliente2 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente2 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente2 + '-' + var_EstadoCliente2 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente2 + '')
    paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente3 + '\n'
                                                                    'Endereço: ' + var_EndCliente3 + ' N°' + var_NCliente3 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente3 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente3 + '-' + var_EstadoCliente3 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente3 + '')
    paragraph.style = document.styles.add_style('style10.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente4 + '\n'
                                                                    'Endereço: ' + var_EndCliente4 + ' N°' + var_NCliente4 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente4 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente4 + '-' + var_EstadoCliente4 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente4 + '')
    paragraph.style = document.styles.add_style('style10.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra:' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Requerimento sem Lei ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '.docx')

    # RRC com lei
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
    paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('\n\nNós,')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner = paragraph.add_run(
        ' ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '')
    runner.bold = True
    paragraph.add_run(
        'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

    paragraph = document.add_paragraph(
        '                                                                         Nestes Termos,\n'
        '                                                                         Pede Deferimento.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dados Complementares:')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Dos Proprietários')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente1 + '\n'
                                                                    'Endereço: ' + var_EndCliente1 + ' N°' + var_NCliente1 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente1 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente1 + '-' + var_EstadoCliente1 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente1 + '')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente2 + '\n'
                                                                    'Endereço: ' + var_EndCliente2 + ' N°' + var_NCliente2 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente2 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente2 + '-' + var_EstadoCliente2 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente2 + '')
    paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente3 + '\n'
                                                                    'Endereço: ' + var_EndCliente3 + ' N°' + var_NCliente3 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente3 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente3 + '-' + var_EstadoCliente3 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente3 + '')
    paragraph.style = document.styles.add_style('style10.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Nome:' + var_NomeCliente4 + '\n'
                                                                    'Endereço: ' + var_EndCliente4 + ' N°' + var_NCliente4 + '\n'
                                                                                                                             'Loteamento:' + var_LoteObra + '\n'
                                                                                                                                                            'CEP:' + var_CEPCliente4 + '\n'
                                                                                                                                                                                       'Cidade/Estado:' + var_CidadeCliente4 + '-' + var_EstadoCliente4 + '\n'
                                                                                                                                                                                                                                                          'Telefone: ' + var_CelCliente4 + '')
    paragraph.style = document.styles.add_style('style10.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('Da Obra')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Endereço: ' + var_EndObra + ' nº ' + var_NObra +
                                       ') – LOTE N° ' + var_NObra + '\n'
                                                                    'Loteamento:' + var_LoteObra + '\n'
                                                                                                   'Quadra:' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                       'CPF: 183.125.858-77\n'
                                       'Celular: (19) 982009858\n'
                                       'Inscrição SMPUGE: 1036/18\n'
                                       'E-mail: rocha.soares@hotmail.com\n')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Requerimento com Lei ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '.docx')

    # Procuração
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('PROCURAÇÃO ')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(11.5)

    runner = paragraph.add_run('I - OUTORGANTES:')
    runner.bold = True

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente1 + ' CPF: ' + var_CPFCliente1 + '')

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente2 + ' CPF: ' + var_CPFCliente2 + '')

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente3 + ' CPF: ' + var_CPFCliente3 + '')

    paragraph.add_run('\nSr.(a) ' + var_NomeCliente4 + ' CPF: ' + var_CPFCliente4 + '\n\n')

    runner = paragraph.add_run('II – OUTORGADO: ')
    runner.bold = True

    paragraph.add_run(
        '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

    runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
    runner.bold = True

    paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                      'Lote' + var_LoteObra + 'da Quadra ' + var_QuadraObra + ', localizado no endereço: '
                                                                              '' + var_EndObra + ' nº ' + var_NObra + ' Loteamento: ' + var_BairroObra + '.\n\n')

    runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
    runner.bold = True

    paragraph.add_run('\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                      '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

    runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + var_CidadeObra + ',')
    runner.bold = True

    paragraph.add_run('especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        'OUTORGANTE:                                                                                                OUTORGANTE:')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente1 + '                                                           ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente1 + '                                                              CPF: ' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n\n\n\n\n_________________________________                            _________________________________')
    paragraph.style = document.styles.add_style('style4.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph(
        'OUTORGANTE:                                                                                                OUTORGANTE:')
    paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph(
        '' + var_NomeCliente3 + '                                                           ' + var_NomeCliente4 + '')
    paragraph.style = document.styles.add_style('style6.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF: ' + var_CPFCliente3 + '                                                               CPF: ' + var_CPFCliente4 + '')
    paragraph.style = document.styles.add_style('style7.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.bold = True

    paragraph = document.add_paragraph('OUTORGADO:')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    font.bold = True

    paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)

    paragraph = document.add_paragraph(
        'CPF: 183.125.858-77')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Procuração ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '.docx')

    # Declaração
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(3.49)
        section.bottom_margin = Cm(1.1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

    paragraph = document.add_paragraph('ANEXO I')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('DECLARAÇÃO')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                       'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                       ' que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município'
                                       ' de Hortolândia”, nós, ' + var_NomeCliente1 + ', (' + var_ProfissaoCliente1 + '), ' + var_NomeCliente2 + ', (' + var_ProfissaoCliente2 + '), ' + var_NomeCliente3 + ', (' + var_ProfissaoCliente3 + ') e ' + var_NomeCliente4 + ', (' + var_ProfissaoCliente4 + '),')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(' Proprietários da obra ')
    runner.bold = True

    paragraph.add_run(
        'localizada à ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ',' \
                                                                                                                                                          'cidade de Hortolândia-SP, DECLARAMOS estar ciente das disposições ' \
                                                                                                                                                          'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                          'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                          'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                          'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                          'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                          'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                          ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                          'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                          'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                          'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                          ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

    runner = paragraph.add_run('DECLARAMOS ')
    runner.bold = True

    paragraph.add_run(
        'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

    runner = paragraph.add_run('um dos seguintes documentos: ')
    runner.bold = True

    paragraph = document.add_paragraph(
        '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run('aaaaaaaaaaa')
    font = runner.font
    font.color.rgb = RGBColor(255, 255, 255)

    paragraph.add_run('origem nativa;')

    paragraph = document.add_paragraph(
        '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente1 + '')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style9.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente3 + '')
    paragraph.style = document.styles.add_style('style9.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n_____________________________________')
    paragraph.style = document.styles.add_style('style8.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('PROPRIETÁRIO ' + var_NomeCliente4 + '')
    paragraph.style = document.styles.add_style('style9.4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('\n\n\n\n\n\n\n\nEm conformidade com o disposto no artigo 4º da '
                                       'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                       'que “Dispõe sobre controle ambiental para utilização'
                                       ' de produtos e subprodutos de madeira de origem nativa'
                                       ' em obras e serviços de Engenharia Civil no Município '
                                       'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                       'Autor do Projeto da obra localizada à '
                                       'Rua ' + var_EndObra + ', nº ' + var_NObra + ' Lote ' + var_LoteObra + ', Quadra ' + var_QuadraObra + ', Loteamento ' + var_BairroObra + ',cidade de Hortolândia-SP,'
                                                                                                                                                                                ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                                                ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                                                'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    runner = paragraph.add_run(
        'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
    runner.underline = True

    paragraph.add_run(
        'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

    paragraph = document.add_paragraph('\n' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = document.add_paragraph('AUTOR DO PROJETO')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.bold = True
    font.size = Pt(10)
    font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # runner_word.size = Pt(10)

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Declaração ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ', ' + var_NomeCliente4 + '.docx')

    # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Obra: REGULARIZAÇÃO E CONSTRUÇÃO RESIDENCIAL MULTIFAMILIAR – R2')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('(DE ACORDO COM A LEI 3.491/2018 - ANISTIA)')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('Local: ' + var_EndObra + '- N° ' + var_NObra + '')
    paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Lote: ' + var_LoteObra + ' Quadra: ' + var_QuadraObra + '')
    paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Loteamento: ' + var_LoteObra + '')
    paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Município: ' + var_CidadeObra + '/SP')
    paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph(
        'Proprietário(s): ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '')
    paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)

    paragraph = document.add_paragraph(
        '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
    paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
    paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
    paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
    paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
    paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
    paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
    paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
    paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
    paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
    paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
    paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
    paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
    paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
    paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
    paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
    paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph(
        '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
    paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph('' + var_CidadeObra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
    paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'
    font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph(
        '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
    paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário: ' + var_NomeCliente1 + '                               Proprietário: ' + var_NomeCliente2 + '')
    paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente1 + '                                                        CPF:' + var_CPFCliente2 + '')
    paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n_________________________________                    _____________________________________')
    paragraph.style = document.styles.add_style('style28.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Proprietário: ' + var_NomeCliente3 + '                               Proprietário: ' + var_NomeCliente4 + '')
    paragraph.style = document.styles.add_style('style29.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CPF:' + var_CPFCliente3 + '                                                        CPF:' + var_CPFCliente4 + '')
    paragraph.style = document.styles.add_style('style30.3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        '\n\n\n\n\n_________________________________')
    paragraph.style = document.styles.add_style('style28.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'ROGÉRIO ROCHA SOARES')
    paragraph.style = document.styles.add_style('style29.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'Engenheiro Civil')
    paragraph.style = document.styles.add_style('style30.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'CREA: 5070347192-SP')
    paragraph.style = document.styles.add_style('style31.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    paragraph = document.add_paragraph(
        'SMPUGE: 1036/18')
    paragraph.style = document.styles.add_style('style32.2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Arial'

    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Memorial Descritivo para Construção ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '.docx')

    # ---------------------Recibo---------------------------------------------------------------------------------------------
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(-4.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.75)

    section = document.sections[0]

    header = document.sections[0].header
    logo = header.paragraphs[0]
    logo_run = logo.add_run()
    logo_run.add_picture("image/logo.png", width=Cm(2.65), height=Cm(2.65))

    paragraph = document.add_paragraph('RECIBO')
    paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                       '5070374192, recebi de ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + ', '
                                                                                                                                                                    'parte do pagamento para aprovação de projeto '
                                                                                                                                                                    'arquitetônico a quantia de R$ ' + var_ValorParcelaContrato + ',00 (' + var_ValorParcelaContratoextenso +
                                       '), de um total de '
                                       'R$ ' + var_ValorContrato + ',00 (' + var_ValorContratoextenso + ').')
    paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.bold = True
    font.name = 'Book Antiqua'
    font.size = Pt(12)

    paragraph = document.add_paragraph('\n\n' + var_CidadeObra + ' ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
    paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = 'Book Antiqua'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font.size = Pt(12)

    # footer section
    footer_section = document.sections[0]
    footer = footer_section.footer

    # footer text
    footer_text = footer.paragraphs[0]
    footer_text.text = "_______________________________________________________________________________________________" \
                       "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                       "\nCREA: 5070347192" \
                       "\nE-MAIL: rocha.soares@hotmail.com"
    document.save(
        './PROCESSO DE CLIENTES/' + var_CidadeObra + '/' + ano + '/' + var_TipoObra + '/' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ' e ' + var_NomeCliente4 + '/Recibo ' + var_NomeCliente1 + ', ' + var_NomeCliente2 + ', ' + var_NomeCliente3 + ', ' + var_NomeCliente4 + '.docx')

    MsgBox = tk.messagebox.showwarning("SUCESSO", "DOCUMENTOS FEITO COM SUCESSO. NÃO ESQUEÇA DE PEGA-LOS",
                                       icon='question')
    if MsgBox == 'ok':
        root.deiconify()
        janela2.destroy()
        JanelaCliente4.destroy()


# Comando para voltar do Cliente 1 para Obra
def volta():
    janela2.deiconify()
    JanelaCliente1.withdraw()

    bntvoltac1 = tk.Button(janela2, text="Voltar Cliente")
    bntvoltac1["command"] = voltac1
    bntvoltac1["width"] = 13
    bntvoltac1.grid(row=19, column=2)

    cb_cliente.config(state='disable')


# Comando para esconder cliente 1 e aparecer obra
def voltac1():
    JanelaCliente1.deiconify()
    janela2.withdraw()


# Comando para voltar do Cliente 2 para Obra
def volta2():
    janela2.deiconify()
    JanelaCliente2.withdraw()

    bntvoltac2 = tk.Button(janela2, text="Voltar Cliente")
    bntvoltac2["command"] = voltac2
    bntvoltac2["width"] = 13
    bntvoltac2.grid(row=17, column=2)

    cb_cliente.config(state='disable')


# Comando para esconder cliente 2 e aparecer obra
def voltac2():
    JanelaCliente2.deiconify()
    janela2.withdraw()


# Comando para voltar do Cliente 3 para Obra
def volta3():
    janela2.deiconify()
    JanelaCliente3.withdraw()

    bntvoltac3 = tk.Button(janela2, text="Voltar Cliente")
    bntvoltac3["command"] = voltac3
    bntvoltac3["width"] = 13
    bntvoltac3.grid(row=17, column=2)

    cb_cliente.config(state='disable')


# Comando para esconder cliente 3 e aparecer obra
def voltac3():
    JanelaCliente3.deiconify()
    janela2.withdraw()


# Comando para voltar do Cliente 4 para Obra
def volta4():
    janela2.deiconify()
    JanelaCliente4.withdraw()

    bntvoltac4 = tk.Button(janela2, text="Voltar Cliente")
    bntvoltac4["command"] = voltac4
    bntvoltac4["width"] = 13
    bntvoltac4.grid(row=17, column=2)

    cb_cliente.config(state='disable')


# Comando para esconder cliente 4 e aparecer obra
def voltac4():
    JanelaCliente4.deiconify()
    janela2.withdraw()


# Copiar cliente 1 para cliente 2
def copiar21():
    EndCliente2.delete(0, END)
    EndCliente2.insert(INSERT, EndCliente1.get())

    NCliente2.delete(0, END)
    NCliente2.insert(INSERT, NCliente1.get())

    CidadeCliente2.delete(0, END)
    CidadeCliente2.insert(INSERT, CidadeCliente1.get())

    BairroCliente2.delete(0, END)
    BairroCliente2.insert(INSERT, BairroCliente1.get())

    CEPCliente2.delete(0, END)
    CEPCliente2.insert(INSERT, CEPCliente1.get())

    EstadoCliente2.delete(0, END)
    EstadoCliente2.insert(INSERT, EstadoCliente1.get())

    cb_EstCivilCliente2.delete(0, END)
    cb_EstCivilCliente2.insert(INSERT, cb_EstCivilCliente1.get())


# Copiar cliente 2 para cliente 3
def copiar32():
    EndCliente3.delete(0, END)
    EndCliente3.insert(INSERT, EndCliente2.get())

    NCliente3.delete(0, END)
    NCliente3.insert(INSERT, NCliente2.get())

    CidadeCliente3.delete(0, END)
    CidadeCliente3.insert(INSERT, CidadeCliente2.get())

    BairroCliente3.delete(0, END)
    BairroCliente3.insert(INSERT, BairroCliente2.get())

    CEPCliente3.delete(0, END)
    CEPCliente3.insert(INSERT, CEPCliente2.get())

    EstadoCliente3.delete(0, END)
    EstadoCliente3.insert(INSERT, EstadoCliente2.get())

    cb_EstCivilCliente3.delete(0, END)
    cb_EstCivilCliente3.insert(INSERT, cb_EstCivilCliente2.get())


# Copiar cliente 3 para cliente 4
def copiar43():
    EndCliente4.delete(0, END)
    EndCliente4.insert(INSERT, EndCliente3.get())

    NCliente4.delete(0, END)
    NCliente4.insert(INSERT, NCliente3.get())

    CidadeCliente4.delete(0, END)
    CidadeCliente4.insert(INSERT, CidadeCliente3.get())

    BairroCliente4.delete(0, END)
    BairroCliente4.insert(INSERT, BairroCliente3.get())

    CEPCliente4.delete(0, END)
    CEPCliente4.insert(INSERT, CEPCliente3.get())

    EstadoCliente4.delete(0, END)
    EstadoCliente4.insert(INSERT, EstadoCliente3.get())

    cb_EstCivilCliente4.delete(0, END)
    cb_EstCivilCliente4.insert(INSERT, cb_EstCivilCliente3.get())


def extenso1():
    ValorContratoextenso.delete(0, END)
    extensoreal1 = (num2words(ValorContrato.get(), lang='pt-br'))
    ValorContratoextenso.insert(INSERT, (extensoreal1 + " reais"))


# Coloca o numero em extenso
def extenso2():
    ParcelContratoextenso.delete(0, END)
    extensoreal2 = (num2words(ParcelContrato.get(), lang='pt-br'))
    ParcelContratoextenso.insert(INSERT, (extensoreal2))


def extenso3():
    ValorParcelaContratoextenso.delete(0, END)
    extensoreal3 = (num2words(ValorParcelaContrato.get(), lang='pt-br'))
    ValorParcelaContratoextenso.insert(INSERT, (extensoreal3 + " reais"))


def extenso4():
    Visitaextenso.delete(0, END)
    extensoreal4 = (num2words(Visita.get(), lang='pt-br'))
    Visitaextenso.insert(INSERT, (extensoreal4 + " reais"))


# Formatação para data
def format_date1(event=None):
    text = InicioContrato.get().replace("/", "").replace("/", "")[:8]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [1, 3]:
            new_text += text[index] + "/"
        else:
            new_text += text[index]

    InicioContrato.delete(0, "end")
    InicioContrato.insert(0, new_text)


# Formatação para Numero
def format_numero1(event=None):
    text = NumObra.get().replace(".", "").replace("-", "")[:50]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [0]:
            new_text += text[index]
        else:
            new_text += text[index]

    NumObra.delete(0, "end")
    NumObra.insert(0, new_text)


# Formatação para CPF do cliente 1
def format_cpf1(event=None):
    text = CPFCliente1.get().replace(".", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [2, 5]:
            new_text += text[index] + "."
        elif index == 8:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CPFCliente1.delete(0, "end")
    CPFCliente1.insert(0, new_text)


# Formatação para RG do cliente 1
def format_rg1(event=None):
    text = RGCliente1.get().replace(".", "").replace("-", "")[:9]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [1, 4]:
            new_text += text[index] + "."
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    RGCliente1.delete(0, "end")
    RGCliente1.insert(0, new_text)


# Formatação para CEP do cliente 1
def format_CEP1(event=None):
    text = CEPCliente1.get().replace("-", "")[:8]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [4]:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CEPCliente1.delete(0, "end")
    CEPCliente1.insert(0, new_text)


# Formatação para Celular do cliente 1
def format_celular1(event=None):
    text = CelCliente1.get().replace("(", "").replace(")", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [0]:
            new_text += "(" + text[index]
        elif index == 1:
            new_text += text[index] + ")"
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CelCliente1.delete(0, "end")
    CelCliente1.insert(0, new_text)


# Formatação para CPF do cliente 2
def format_cpf2(event=None):
    text = CPFCliente2.get().replace(".", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [2, 5]:
            new_text += text[index] + "."
        elif index == 8:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CPFCliente2.delete(0, "end")
    CPFCliente2.insert(0, new_text)


# Formatação para RG do cliente 2
def format_rg2(event=None):
    text = RGCliente2.get().replace(".", "").replace("-", "")[:9]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [1, 4]:
            new_text += text[index] + "."
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    RGCliente2.delete(0, "end")
    RGCliente2.insert(0, new_text)


# Formatação para CEP do cliente 2
def format_CEP2(event=None):
    text = CEPCliente2.get().replace("-", "")[:8]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [4]:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CEPCliente2.delete(0, "end")
    CEPCliente2.insert(0, new_text)


# Formatação para Celular do cliente 2
def format_celular2(event=None):
    text = CelCliente2.get().replace("(", "").replace(")", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [0]:
            new_text += "(" + text[index]
        elif index == 1:
            new_text += text[index] + ")"
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CelCliente2.delete(0, "end")
    CelCliente2.insert(0, new_text)


# Formatação para CPF do cliente 3
def format_cpf3(event=None):
    text = CPFCliente3.get().replace(".", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [2, 5]:
            new_text += text[index] + "."
        elif index == 8:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CPFCliente3.delete(0, "end")
    CPFCliente3.insert(0, new_text)


# Formatação para RG do cliente 3
def format_rg3(event=None):
    text = RGCliente3.get().replace(".", "").replace("-", "")[:9]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [1, 4]:
            new_text += text[index] + "."
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    RGCliente3.delete(0, "end")
    RGCliente3.insert(0, new_text)


# Formatação para CEP do cliente 3
def format_CEP3(event=None):
    text = CEPCliente3.get().replace("-", "")[:8]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [4]:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CEPCliente3.delete(0, "end")
    CEPCliente3.insert(0, new_text)


# Formatação para Celular do cliente 3
def format_celular3(event=None):
    text = CelCliente3.get().replace("(", "").replace(")", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [0]:
            new_text += "(" + text[index]
        elif index == 1:
            new_text += text[index] + ")"
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CelCliente3.delete(0, "end")
    CelCliente3.insert(0, new_text)


# Formatação para CPF do cliente 4
def format_cpf4(event=None):
    text = CPFCliente4.get().replace(".", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [2, 5]:
            new_text += text[index] + "."
        elif index == 8:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CPFCliente4.delete(0, "end")
    CPFCliente4.insert(0, new_text)


# Formatação para RG do cliente 4
def format_rg4(event=None):
    text = RGCliente4.get().replace(".", "").replace("-", "")[:9]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [1, 4]:
            new_text += text[index] + "."
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    RGCliente4.delete(0, "end")
    RGCliente4.insert(0, new_text)


# Formatação para CEP do cliente 4
def format_CEP4(event=None):
    text = CEPCliente4.get().replace("-", "")[:8]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [4]:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CEPCliente4.delete(0, "end")
    CEPCliente4.insert(0, new_text)


# Formatação para Celular do cliente 4
def format_celular4(event=None):
    text = CelCliente4.get().replace("(", "").replace(")", "").replace("-", "")[:11]
    new_text = ""

    if event.keysym.lower() == "backspace": return

    for index in range(len(text)):

        if not text[index] in "0123456789": continue
        if index in [0]:
            new_text += "(" + text[index]
        elif index == 1:
            new_text += text[index] + ")"
        elif index == 6:
            new_text += text[index] + "-"
        else:
            new_text += text[index]

    CelCliente4.delete(0, "end")
    CelCliente4.insert(0, new_text)


img = PhotoImage(file="image/logo.png")

imagen = Label(image=img).pack()

root.configure(bg='white')

root.bind("<KeyPress>", lambda e: Next() if e.char == '\r' else None)

root.bind('<Button-1>', lambda e: Next())

root.config(cursor="pirate")

registro = Label(root, text="Open Spurce - Syller", font='underline')
registro.pack(side=BOTTOM)

root.mainloop()

